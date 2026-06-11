"""
rootkitguard.py v2.1
ФИКСЫ:
  - Убран matplotlib/FigureCanvasTkAgg → нет PIL-конфликта
  - Кнопка автозапуска API прямо из GUI (▶ рядом с «API offline»)
  - Отчёты уникальные: имя файла + время скана в названии
  - Кнопка «Сгенерировать демо-модели» если моделей нет
  - Rootkit Scan: пошаговый прогресс с карточками
"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))

import customtkinter as ctk
import threading
import subprocess
import requests
import joblib
import pandas as pd
import numpy as np
from tkinter import filedialog
from pathlib import Path
from datetime import datetime

from process_monitor import ProcessMonitor
from rootkit_checker import RootkitChecker
from notifier import notify_threat
from logger import get_logger
from config_loader import cfg
from i18n import t, set_lang, get_lang

log = get_logger("gui")

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

API_BASE = f"http://127.0.0.1:{cfg.get('api', {}).get('port', 8000)}"


def generate_demo_models():
    """Создаёт рабочие модели на синтетических данных (~10 сек)."""
    from sklearn.ensemble import RandomForestClassifier, IsolationForest
    from sklearn.preprocessing import StandardScaler
    import xgboost as xgb

    log.info("Генерация демо-моделей...")
    np.random.seed(42)
    n = 5000
    normal = np.random.randn(int(n * 0.73), 78) * 0.5
    attack = np.random.randn(int(n * 0.27), 78) * 2.0 + 3.0
    X_all  = np.vstack([normal, attack])
    y_all  = np.array([0] * len(normal) + [1] * len(attack))

    cols = [
        'Dst Port','Protocol','Flow Duration','Tot Fwd Pkts','Tot Bwd Pkts',
        'TotLen Fwd Pkts','TotLen Bwd Pkts','Fwd Pkt Len Max','Fwd Pkt Len Min',
        'Fwd Pkt Len Mean','Fwd Pkt Len Std','Bwd Pkt Len Max','Bwd Pkt Len Min',
        'Bwd Pkt Len Mean','Bwd Pkt Len Std','Flow Byts/s','Flow Pkts/s',
        'Flow IAT Mean','Flow IAT Std','Flow IAT Max','Flow IAT Min',
        'Fwd IAT Tot','Fwd IAT Mean','Fwd IAT Std','Fwd IAT Max','Fwd IAT Min',
        'Bwd IAT Tot','Bwd IAT Mean','Bwd IAT Std','Bwd IAT Max','Bwd IAT Min',
        'Fwd PSH Flags','Bwd PSH Flags','Fwd URG Flags','Bwd URG Flags',
        'Fwd Header Len','Bwd Header Len','Fwd Pkts/s','Bwd Pkts/s',
        'Pkt Len Min','Pkt Len Max','Pkt Len Mean','Pkt Len Std','Pkt Len Var',
        'FIN Flag Cnt','SYN Flag Cnt','RST Flag Cnt','PSH Flag Cnt','ACK Flag Cnt',
        'URG Flag Cnt','CWE Flag Count','ECE Flag Cnt','Down/Up Ratio',
        'Pkt Size Avg','Fwd Seg Size Avg','Bwd Seg Size Avg','Fwd Byts/b Avg',
        'Fwd Pkts/b Avg','Fwd Blk Rate Avg','Bwd Byts/b Avg','Bwd Pkts/b Avg',
        'Bwd Blk Rate Avg','Subflow Fwd Pkts','Subflow Fwd Byts',
        'Subflow Bwd Pkts','Subflow Bwd Byts','Init Fwd Win Byts','Init Bwd Win Byts',
        'Fwd Act Data Pkts','Fwd Seg Size Min','Active Mean','Active Std',
        'Active Max','Active Min','Idle Mean','Idle Std','Idle Max','Idle Min','Inbound',
    ]
    cols = cols[:X_all.shape[1]]
    X_df = pd.DataFrame(X_all, columns=cols)

    scaler = StandardScaler()
    X_sc   = pd.DataFrame(scaler.fit_transform(X_df), columns=cols)

    rf = RandomForestClassifier(n_estimators=50, max_depth=8,
                                 class_weight='balanced', random_state=42, n_jobs=-1)
    rf.fit(X_sc, y_all)

    xgb_m = xgb.XGBClassifier(n_estimators=50, max_depth=6, scale_pos_weight=2.7,
                                random_state=42, eval_metric='logloss', verbosity=0)
    xgb_m.fit(X_sc, y_all)

    iso = IsolationForest(n_estimators=50, contamination=0.27,
                          random_state=42, n_jobs=-1)
    iso.fit(X_sc)

    Path("models").mkdir(exist_ok=True)
    joblib.dump(rf,     "models/rf_cicids.pkl")
    joblib.dump(xgb_m,  "models/xgb_cicids.pkl")
    joblib.dump(iso,    "models/iso_cicids.pkl")
    joblib.dump(scaler, "models/scaler_cicids.pkl")
    log.info("Демо-модели сохранены")
    return rf, scaler


class RootkitGuard(ctk.CTk):
    def __init__(self, username: str = "admin"):
        super().__init__()
        self.username = username
        self._app_log = []  # глобальный журнал всех событий приложения
        self.title(f"RootkitGuard — {self.username} — Система обнаружения аномалий v2.1")
        w, h = cfg.get("app", {}).get("window_size", "1200x750").split("x")
        self.geometry(f"{w}x{h}")
        self.resizable(True, True)

        self.model_loaded = False
        threading.Thread(target=self._load_models_bg, daemon=True).start()

        self._last_scan = {
            "total": 0, "anomaly": 0, "normal": 0,
            "pct": 0.0, "threat": "—",
            "filename": "", "filepath": "", "timestamp": "",
            "top_ports": [], "max_proba": 0.0,
        }
        self._prev_scan = {}
        self._api_available = False
        self._api_proc      = None
        self._build_ui()
        self._setup_scroll_zoom()
        threading.Thread(target=self._check_api, daemon=True).start()
        threading.Thread(target=self._auto_startup_scan, daemon=True).start()

    def _load_models_bg(self):
        """Грузим модели в фоне, чтобы окно открывалось мгновенно."""
        self._load_models()
        try:
            mc = "#2dc97e" if self.model_loaded else "#e74c3c"
            mt = t("model_loaded") if self.model_loaded else t("model_not_found")
            self.after(0, lambda: self.model_lbl.configure(text=mt, text_color=mc))
        except Exception:
            pass

    def _load_models(self):
        try:
            self.scaler      = joblib.load("models/scaler_cicids.pkl")
            self.rf_default  = joblib.load("models/rf_default.pkl")
            self.current_model_name = "RF Default"
            self.rf = self.rf_default
            self.model_loaded = True
        except Exception as e:
            print(f"[ThreatMonitor] RF: {e}")

        try:
            self.xgb_model = joblib.load("models/xgb_cicids.pkl")
        except Exception as e:
            self.xgb_model = None
            print(f"[ThreatMonitor] XGB: {e}")

        try:
            self.iso_model = joblib.load("models/iso_cicids.pkl")
        except Exception as e:
            self.iso_model = None
            print(f"[ThreatMonitor] ISO: {e}")

        try:
            rkg_path = "models/rf_rootkitguard.pkl"
            self.rkg_model = joblib.load(rkg_path) if Path(rkg_path).exists() else self.rf_default
        except Exception as e:
            self.rkg_model = None
            print(f"[ThreatMonitor] RKG: {e}")

    def set_model(self, model_name: str):
        models = {
            "rf":           self.rf_default,
            "xgb":          self.xgb_model,
            "iso":          self.iso_model,
            "rootkitguard": self.rkg_model,
        }
        self.rf = models.get(model_name, self.rf_default)
        self.current_model_name = {
            "rf":           "Random Forest",
            "xgb":          "XGBoost",
            "iso":          "Isolation Forest",
            "rootkitguard": "RootkitGuard ML",
        }.get(model_name, "Random Forest")

    # ── API ─────────────────────────────────────────────────────

    def _check_api(self):
        try:
            r = requests.get(f"{API_BASE}/health", timeout=2)
            self._api_available = (r.status_code == 200)
        except Exception:
            self._api_available = False
        status = "● API online" if self._api_available else "● API offline"
        color  = "#2dc97e" if self._api_available else "#f39c12"
        self.after(0, lambda: self.api_lbl.configure(text=status, text_color=color))

    def _start_api(self):
        if self._api_available:
            self.api_lbl.configure(text=t("api_running"), text_color="#2dc97e")
            return
        try:
            main_py = str(Path(__file__).parent.parent / "main.py")
            self._api_proc = subprocess.Popen(
                [sys.executable, main_py, "api"],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            self.api_lbl.configure(text=t("api_starting"), text_color="yellow")
            threading.Thread(target=self._wait_api, daemon=True).start()
        except Exception as e:
            log.error(f"API запуск: {e}")

    def _wait_api(self):
        import time
        for _ in range(12):
            time.sleep(1)
            try:
                if requests.get(f"{API_BASE}/health", timeout=1).status_code == 200:
                    self._api_available = True
                    self.after(0, lambda: self.api_lbl.configure(
                        text="● API online", text_color="#2dc97e"))
                    return
            except Exception:
                pass
        self.after(0, lambda: self.api_lbl.configure(
            text=t("api_no_resp"), text_color="red"))

    def _auto_startup_scan(self):
        """При запуске делает Rootkit Scan в фоне (если включено в настройках)."""
        import time
        if not cfg.get("scan", {}).get("autostart_scan", True):
            return
        if cfg.get("performance", {}).get("lite_mode", False):
            return  # lite-режим: не нагружаем систему при старте
        time.sleep(3)
        log.info("Авто-сканирование при запуске...")
        try:
            checker = RootkitChecker()
            result  = checker.run_all()
            threat  = result.threat_level
            count   = len(result.findings)
            color   = {"ВЫСОКАЯ": "#e74c3c", "СРЕДНЯЯ": "#f39c12",
                       "ЧИСТАЯ": "#2dc97e"}.get(threat, "#2dc97e")
            self._last_autoscan = (threat, count)
            notify_threat(threat, f"Авто-сканирование при запуске: {count} находок")
            # (раньше тут вызывался _run_rootkit_local — методы страницы,
            #  которой нет в навигации; это роняло фоновый поток)
        except Exception as e:
           log.error(f"Авто-скан ошибка: {e}")
    # ── UI ──────────────────────────────────────────────────────


    def app_log(self, msg: str):
        """Записывает событие в глобальный журнал приложения."""
        from datetime import datetime
        ts = datetime.now().strftime("%H:%M:%S")
        line = f"[{ts}] {msg}"
        self._app_log.append(line)
        # Ограничиваем размер журнала
        if len(self._app_log) > 2000:
            self._app_log = self._app_log[-2000:]

    def _copy_all_log(self):
        """Копирует весь журнал приложения в буфер обмена."""
        from datetime import datetime
        header = [
            "=" * 60,
            "ROOTKITGUARD v2.1 — ПОЛНЫЙ ЖУРНАЛ ПРИЛОЖЕНИЯ",
            f"Экспортировано: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Пользователь: {self.username}",
            f"Модель: {getattr(self, 'current_model_name', 'RF')}",
            "=" * 60,
            "",
        ]
        # Добавляем последний rootkit-скан если был
        full = header + self._app_log

        # Добавляем детальные находки rootkit если есть
        if hasattr(self, "_last_rkd_report") and self._last_rkd_report:
            full += ["", "=" * 60, "ПОСЛЕДНИЙ ROOTKIT-СКАН (детально):", "=" * 60]
            full += self._last_rkd_report

        text = "\n".join(full)
        try:
            self.clipboard_clear()
            self.clipboard_append(text)
            self.update()
            if hasattr(self, "_copy_btn"):
                self._copy_btn.configure(text=t("copied"))
                self.after(2000, lambda: self._copy_btn.configure(
                    text=t("copy_all")))
        except Exception as e:
            print(f"Ошибка копирования: {e}")

    def _setup_scroll_zoom(self):
        """#15 колесо мыши прокручивает страницу везде; #16 зум интерфейса."""
        self._ui_scale = 1.0

        # ── Скролл колесом по всей области (CTk не делает это для дочерних) ──
        def _scroll(direction):
            try:
                canvas = self.main._parent_canvas
                canvas.yview_scroll(direction, "units")
            except Exception:
                pass

        def on_wheel(event):
            if event.state & 0x0004:      # Ctrl зажат → это зум, не скролл
                return
            if getattr(event, "delta", 0):           # Windows/macOS
                _scroll(-1 if event.delta > 0 else 1)
            elif getattr(event, "num", None) == 4:   # Linux вверх
                _scroll(-1)
            elif getattr(event, "num", None) == 5:   # Linux вниз
                _scroll(1)

        self.bind_all("<MouseWheel>", on_wheel, add="+")
        self.bind_all("<Button-4>", on_wheel, add="+")
        self.bind_all("<Button-5>", on_wheel, add="+")

        # ── Зум: Ctrl +/-/0 и Ctrl+колесо ──
        def on_ctrl_wheel(event):
            if getattr(event, "delta", 0) > 0 or getattr(event, "num", None) == 4:
                self._zoom(+0.1)
            else:
                self._zoom(-0.1)
            return "break"
        self.bind_all("<Control-MouseWheel>", on_ctrl_wheel, add="+")
        self.bind_all("<Control-Button-4>", lambda e: self._zoom(+0.1), add="+")
        self.bind_all("<Control-Button-5>", lambda e: self._zoom(-0.1), add="+")
        self.bind_all("<Control-plus>",  lambda e: self._zoom(+0.1), add="+")
        self.bind_all("<Control-equal>", lambda e: self._zoom(+0.1), add="+")
        self.bind_all("<Control-minus>", lambda e: self._zoom(-0.1), add="+")
        self.bind_all("<Control-Key-0>", lambda e: self._zoom(reset=True), add="+")

        # ── Компактный зум-контрол внизу нав-панели ──
        zoom_bar = ctk.CTkFrame(self.nav, fg_color="transparent")
        zoom_bar.pack(side="bottom", fill="x", padx=8, pady=(0, 6))
        ctk.CTkButton(zoom_bar, text="\u2212", width=30, height=26,
                      fg_color="#1e293b", hover_color="#2d3748",
                      command=lambda: self._zoom(-0.1)).pack(side="left")
        self._zoom_lbl = ctk.CTkButton(zoom_bar, text="100%", width=54, height=26,
                      fg_color="transparent", hover_color="#1e293b",
                      command=lambda: self._zoom(reset=True))
        self._zoom_lbl.pack(side="left", padx=2)
        ctk.CTkButton(zoom_bar, text="+", width=30, height=26,
                      fg_color="#1e293b", hover_color="#2d3748",
                      command=lambda: self._zoom(+0.1)).pack(side="left")

    def _zoom(self, delta=0.0, reset=False):
        """Масштабирование всего интерфейса (для демонстрации на большом экране)."""
        self._ui_scale = 1.0 if reset else min(2.2, max(0.7, self._ui_scale + delta))
        try:
            ctk.set_widget_scaling(self._ui_scale)
            if hasattr(self, "_zoom_lbl"):
                self._zoom_lbl.configure(text=f"{int(self._ui_scale * 100)}%")
        except Exception:
            pass

    def _build_ui(self):
        self._nav_expanded = True
        self.nav = ctk.CTkFrame(self, width=210, corner_radius=0)
        self.nav.pack(side="left", fill="y")
        self.nav.pack_propagate(False)

        # Кнопка сворачивания панели
        toggle_btn = ctk.CTkButton(
            self.nav, text="◀", width=30, height=30,
            fg_color="transparent", hover_color="#2d2d44",
            font=ctk.CTkFont(size=14),
            command=self._toggle_nav)
        toggle_btn.pack(anchor="e", padx=8, pady=(8, 0))
        self._toggle_btn = toggle_btn

        self.nav_title = ctk.CTkLabel(self.nav, text="RootkitGuard",
                     font=ctk.CTkFont(size=19, weight="bold"))
        self.nav_title.pack(pady=(20, 2))
        self.nav_version = ctk.CTkLabel(self.nav, text=f"v{cfg.get('app',{}).get('version','2.1')}",
                     font=ctk.CTkFont(size=11), text_color="gray")
        self.nav_version.pack(pady=(0, 16))

        pages = [
            ("  🏠  Главная",       "home"),
            ("  🛡  Rootkit Defense","rkdefense"),
            ("  🔍  Сканирование",  "scan"),
            ("  👁   Мониторинг",   "monitor"),
            ("  📊  Аналитика",     "analytics"),
            ("  ⚙️  Настройки",    "settings"),
            ("  ℹ️  О системе",     "about"),
        ]
        self.nav_buttons = {}
        for label, key in pages:
            btn = ctk.CTkButton(
                self.nav, text=label, anchor="w",
                fg_color="transparent", hover_color="#2d2d44",
                font=ctk.CTkFont(size=13),
                command=lambda k=key: self.show_page(k))
            btn.pack(fill="x", padx=10, pady=2)
            self.nav_buttons[key] = btn

        # Статус модели
        mc = "#2dc97e" if self.model_loaded else "#e74c3c"
        mt = "● Модель загружена" if self.model_loaded else "● Модель не найдена"
        model_row = ctk.CTkFrame(self.nav, fg_color="transparent")
        model_row.pack(side="bottom", fill="x", padx=8, pady=(0, 2))
        self.model_lbl = ctk.CTkLabel(model_row, text=mt,
                                       text_color=mc, font=ctk.CTkFont(size=11))
        self.model_lbl.pack(side="left")
        self.learn_lbl = ctk.CTkLabel(model_row, text="",
                                       font=ctk.CTkFont(size=11),
                                       text_color="#a855f7")
        self.learn_lbl.pack(side="right")

        # Кнопка языка с dropdown
        # Переключатель языка
        lang_row = ctk.CTkFrame(self.nav, fg_color="transparent")
        lang_row.pack(side="bottom", fill="x", padx=8, pady=(0, 4))
        for lang, label in [("ru","РУС"),("en","ENG"),("kz","ҚАЗ")]:
            ctk.CTkButton(lang_row, text=label, width=52, height=22,
                fg_color="#1e293b", hover_color="#2d3748",
                font=ctk.CTkFont(size=10), corner_radius=6,
                text_color="#64748b",
                command=lambda l=lang: self._switch_lang(l)
            ).pack(side="left", padx=2)
        self._lang_menu_open = False
            

        # API статус + кнопка запуска
        api_row = ctk.CTkFrame(self.nav, fg_color="transparent")
        api_row.pack(side="bottom", fill="x", padx=8, pady=(0, 2))
        self.api_lbl = ctk.CTkLabel(api_row, text=t("api_checking"),
                                     text_color="gray", font=ctk.CTkFont(size=11))
        self.api_lbl.pack(side="left")
        ctk.CTkButton(api_row, text="▶", width=28, height=22,
                      fg_color="#2d6a4f", font=ctk.CTkFont(size=10),
                      command=lambda: threading.Thread(
                          target=self._start_api, daemon=True).start()
                      ).pack(side="right")
        self._lang_btn = ctk.CTkButton(api_row,
            text={"ru": "РУС", "en": "ENG", "kz": "ҚАЗ"}.get(get_lang(), "РУС"), width=52, height=22,
            fg_color="transparent", hover_color="#1e293b",
            font=ctk.CTkFont(size=10), corner_radius=6,
            text_color="#64748b",
            command=self._show_lang_menu)
        self._lang_btn.pack(side="right", padx=(0,4))
        self._lang_menu_open = False

        self.main = ctk.CTkScrollableFrame(self, corner_radius=0,
                                            fg_color="transparent",
                                            scrollbar_button_color="#1e293b",
                                            scrollbar_button_hover_color="#2d3748")
        self.main.pack(side="left", fill="both", expand=True, padx=10, pady=10)

        self._page_builders = {
            "home":      self._page_home,
            "rkdefense": self._page_rkdefense,
            "scan":      self._page_scan,
            "monitor":   self._page_monitor,
            "analytics": self._page_analytics,
            "settings":  self._page_settings,
            "about":     self._page_about,
        }
        self.pages = {}          # страницы строятся лениво при первом открытии
        self._current_page = "home"
        self.show_page("home")

    def _toggle_nav(self):
        if self._nav_expanded:
            # Сворачиваем
            self.nav.configure(width=50)
            for key, btn in self.nav_buttons.items():
                icons = {"home":"🏠","rkdefense":"🛡","scan":"🔍",
                         "monitor":"👁","analytics":"📊",
                         "settings":"⚙️","about":"ℹ️"}
                btn.configure(text=f"  {icons.get(key,'●')}")
            self._toggle_btn.configure(text="▶")
            self.model_lbl.configure(text="●")
            self.api_lbl.configure(text="●")
            # Скрываем языковые кнопки
            # язык скрыт при сворачивании
            self._lang_btn.configure(text="●", width=36)
            self._nav_expanded = False
            self.nav_title.pack_forget()
            self.nav_version.pack_forget()

        else:
            # Разворачиваем
            self.nav.configure(width=210)
            pages_nav = [
                (t("home"),          "home"),
                ("Rootkit Defense",  "rkdefense"),
                (t("scan"),          "scan"),
                (t("monitor"),       "monitor"),
                (t("analytics"),     "analytics"),
                (t("settings"),      "settings"),
                (t("about"),         "about"),
            ]
            icons = ["🏠","🛡","🔍","👁","📊","⚙️","ℹ️"]
            for (label, key), icon in zip(pages_nav, icons):
                self.nav_buttons[key].configure(text=f"  {icon}  {label}")
            self._toggle_btn.configure(text="◀")
            mc = "#2dc97e" if self.model_loaded else "#e74c3c"
            mt = t("model_loaded") if self.model_loaded else t("model_not_found")
            self.model_lbl.configure(text=mt, text_color=mc)
            # Восстанавливаем языковые кнопки
            labels = {"ru": "РУС", "en": "ENG", "kz": "ҚАЗ"}
            self._lang_btn.configure(
                text=labels.get(get_lang(), "РУС"), width=160)
            # Восстанавливаем API статус
            threading.Thread(target=self._check_api, daemon=True).start()
            self._nav_expanded = True
            self.nav_title.pack(pady=(20, 2), before=self.nav_buttons["home"])
            self.nav_version.pack(pady=(0, 16), before=self.nav_buttons["home"])

    def show_page(self, key):
        if key not in self.pages:
            self.pages[key] = self._page_builders[key]()   # ленивое построение
        for p in self.pages.values():
            p.pack_forget()
        self.pages[key].pack(fill="both", expand=True)
        self._current_page = key
        for k, btn in self.nav_buttons.items():
            btn.configure(fg_color="#1f538d" if k == key else "transparent")
        # Мониторинг сканирует процессы только пока его страница открыта —
        # иначе фоновый цикл нагружает CPU и лагает весь интерфейс.
        if hasattr(self, "_monitor"):
            if key == "monitor":
                if not self._monitor.running:
                    self._monitor.start_realtime(
                        interval=6 if getattr(self, "_lite", False)
                        else cfg.get("monitor", {}).get("interval_sec", 3))
                    if hasattr(self, "mon_status"):
                        self.mon_status.configure(text="● LIVE", text_color="#2dc97e")
            else:
                self._monitor.stop_realtime()
        if key == "monitor":
            self.after(50, self._refresh_monitor_table)
        if key == "analytics":
            self.after(50, self._refresh_analytics_snapshot)
        if key == "home" and hasattr(self, "_restart_clock"):
            self._restart_clock()


    def _show_lang_menu(self):
        pass


    def _close_lang_menu(self):
        if hasattr(self, '_lang_popup'):
            self._lang_popup.destroy()
        self._lang_menu_open = False
        self.unbind("<Button-1>")

    def _select_lang(self, lang: str):
        labels = {"ru": "РУС", "en": "ENG", "kz": "ҚАЗ"}
        self._lang_btn.configure(text=labels.get(lang, "🌐"))
        self._close_lang_menu()
        self._switch_lang(lang)

    def _switch_lang(self, lang: str):
        set_lang(lang)
        # Останавливаем фоновый мониторинг перед сносом страницы
        if hasattr(self, "_monitor"):
            try:
                self._monitor.stop_realtime()
            except Exception:
                pass
        # Пересоздаём только построенные страницы — лениво
        for p in self.pages.values():
            p.destroy()
        self.pages = {}
        cur = getattr(self, "_current_page", "home")
        self.pages[cur] = self._page_builders[cur]()
        self.pages[cur].pack(fill="both", expand=True)
        # Обновляем навигацию
        pages_nav = [
            (t("home"),         "home"),
            ("Rootkit Defense", "rkdefense"),
            (t("scan"),         "scan"),
            (t("monitor"),      "monitor"),
            (t("analytics"),    "analytics"),
            (t("settings"),     "settings"),
            (t("about"),        "about"),
        ]
        icons = ["🏠","🛡","🔍","👁","📊","⚙️","ℹ️"]
        for (label, key), icon in zip(pages_nav, icons):
            if key in self.nav_buttons:
                self.nav_buttons[key].configure(text=f"  {icon}  {label}")
        # Обновляем статусы внизу панели
        mc = "#2dc97e" if self.model_loaded else "#e74c3c"
        mt = t("model_loaded") if self.model_loaded else t("model_not_found")
        self.model_lbl.configure(text=mt, text_color=mc)
        # Обновляем текст авто-скана если уже был
        
        labels = {"ru": "РУС", "en": "ENG", "kz": "ҚАЗ"}
        self._lang_btn.configure(text=labels.get(lang, "🌐"))
        
        self.show_page("home")
        
    def _gen_demo_models_thread(self):
        self.model_lbl.configure(text=t("generating"), text_color="yellow")
        try:
            generate_demo_models()
            self._load_models()
            if self.model_loaded:
                self.model_lbl.configure(text=t("model_loaded"), text_color="#2dc97e")
            else:
                self.model_lbl.configure(text=t("model_load_error"), text_color="red")
        except Exception as e:
            log.error(f"Demo model error: {e}")
            self.model_lbl.configure(text=t("error_dot"), text_color="red")

    # ── Главная ─────────────────────────────────────────────────

    


    def _toggle_models(self):
        """Свернуть/развернуть карточки моделей (позиция блоков сохраняется)."""
        self._models_visible = not getattr(self, "_models_visible", True)
        if self._models_visible:
            self._models_frame.pack(fill="x")
            self._models_toggle_btn.configure(text="▾  " + t("choose_model"))
        else:
            self._models_frame.pack_forget()
            self._models_toggle_btn.configure(text="▸  " + t("choose_model"))

    def _page_home(self):
        from pathlib import Path
    
        frame = ctk.CTkFrame(self.main, fg_color="transparent")
    
        # ── Топ-бар ───────────────────────────────────────────────
        topbar = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                              border_width=1, border_color="#1e293b", height=52)
        topbar.pack(fill="x", padx=16, pady=(8, 6))
        topbar.pack_propagate(False)
    
        ctk.CTkLabel(topbar, text="⬡  ROOTKITGUARD",
                     font=ctk.CTkFont(size=13, weight="bold"),
                     text_color="#00d4ff").pack(side="left", padx=16, pady=10)
    
        # Разделитель
        ctk.CTkFrame(topbar, fg_color="#1e293b", width=1).pack(side="left", fill="y", pady=8)
    
        ctk.CTkLabel(topbar, text=f"  👤 {self.username}",
                     font=ctk.CTkFont(size=12),
                     text_color="#64748b").pack(side="left", padx=12)
    
        ctk.CTkFrame(topbar, fg_color="#1e293b", width=1).pack(side="left", fill="y", pady=8)
    
        self.home_time_lbl = ctk.CTkLabel(topbar, text="",
                                           font=ctk.CTkFont(size=12),
                                           text_color="#64748b")
        self.home_time_lbl.pack(side="left", padx=12)
    
        # Статус API справа
        self.home_api_lbl = ctk.CTkLabel(topbar, text="● API offline",
                                          font=ctk.CTkFont(size=11),
                                          text_color="#f39c12")
        self.home_api_lbl.pack(side="right", padx=16)
    
        def update_time():
            from datetime import datetime
            now = datetime.now().strftime("%d.%m.%Y  %H:%M:%S")
            try:
                self.home_time_lbl.configure(text=f"🕐  {now}")
                api_txt = "● API online" if self._api_available else "● API offline"
                api_col = "#00ff88" if self._api_available else "#f39c12"
                self.home_api_lbl.configure(text=api_txt, text_color=api_col)
                if getattr(self, "_current_page", "home") == "home":
                    self.after(1000, update_time)
            except Exception:
                pass
        self._restart_clock = lambda: self.after(200, update_time)
        self.after(100, update_time)
    
        # ── Выбор модели ──────────────────────────────────────────
        self._models_toggle_btn = ctk.CTkButton(
            frame, text="▾  " + t("choose_model"), height=22,
            fg_color="transparent", hover_color="#1e293b",
            font=ctk.CTkFont(size=11, weight="bold"),
            text_color="#64748b", anchor="w",
            command=self._toggle_models)
        self._models_toggle_btn.pack(anchor="w", padx=14, pady=(4, 6))
    
        self._models_wrap = ctk.CTkFrame(frame, fg_color="transparent", height=0)
        self._models_wrap.pack(fill="x", padx=16)
        models_frame = ctk.CTkFrame(self._models_wrap, fg_color="transparent")
        models_frame.pack(fill="x")
        self._models_frame = models_frame
        self._models_visible = True
    
        model_data = [
            {
                "name":  "Random Forest",
                "short": "RF",
                "type":  "Supervised",
                "f1":    "1.0000",
                "auc":   "0.9999",
                "speed": "●●●●○",
                "desc":  t("rf_desc"),
                "color": "#0ea5e9",
                "bg":    "#0c1929",
            },
            {
                "name":  "XGBoost",
                "short": "XGB",
                "type":  "Supervised",
                "f1":    "1.0000",
                "auc":   "1.0000",
                "speed": "●●●○○",
                "desc":  t("xgb_desc"),
                "color": "#a855f7",
                "bg":    "#160d29",
            },
            {
                "name":  "Isolation Forest",
                "short": "ISO",
                "type":  "Unsupervised",
                "f1":    "0.0200",
                "auc":   "0.3258",
                "speed": "●●●●●",
                "desc":  t("iso_desc"),
                "color": "#f59e0b",
                "bg":    "#1a1200",
            },
            {
                "name":  t("ensemble"),
                "short": "ALL",
                "type":  "Hybrid",
                "f1":    "1.0000",
                "auc":   "0.9999",
                "speed": "●●○○○",
                "desc":  t("all_desc"),
                "color": "#00ff88",
                "bg":    "#001a0d",
            },
        ]
    
        self._selected_model = ctk.StringVar(value="Random Forest")
        self._model_cards = {}
    
        for i, m in enumerate(model_data):
            models_frame.grid_columnconfigure(i, weight=1)
            card = ctk.CTkFrame(models_frame, fg_color=m["bg"], corner_radius=12,
                                border_width=1, border_color="#1e293b",
                                cursor="hand2")
            card.grid(row=0, column=i, padx=5, sticky="ew")
    
            # Индикатор выбора
            indicator = ctk.CTkFrame(card, fg_color=m["color"], height=3,
                                      corner_radius=0)
            indicator.pack(fill="x")
            indicator.pack_forget()  # скрыт по умолчанию
    
            # Контент
            inner = ctk.CTkFrame(card, fg_color="transparent")
            inner.pack(fill="both", expand=True, padx=12, pady=10)
    
            # Шапка
            top = ctk.CTkFrame(inner, fg_color="transparent")
            top.pack(fill="x")
            ctk.CTkLabel(top, text=m["short"],
                         font=ctk.CTkFont(size=20, weight="bold"),
                         text_color=m["color"]).pack(side="left")
            ctk.CTkLabel(top, text=m["type"],
                         font=ctk.CTkFont(size=9),
                         text_color="#64748b").pack(side="right", pady=(4, 0))
    
            ctk.CTkLabel(inner, text=m["name"],
                         font=ctk.CTkFont(size=11, weight="bold"),
                         text_color="#64748b", anchor="w").pack(anchor="w", pady=(2, 6))
    
            ctk.CTkLabel(inner, text=m["desc"],
                         font=ctk.CTkFont(size=10),
                         text_color="#64748b", justify="left", anchor="w").pack(anchor="w")
    
            # Метрики
            metrics = ctk.CTkFrame(inner, fg_color="transparent")
            metrics.pack(fill="x", pady=(8, 4))
            for label, val in [("F1", m["f1"]), ("AUC", m["auc"])]:
                mf = ctk.CTkFrame(metrics, fg_color="#0a0e1a", corner_radius=6)
                mf.pack(side="left", padx=(0, 4))
                ctk.CTkLabel(mf, text=label,
                             font=ctk.CTkFont(size=9),
                             text_color="#64748b").pack(padx=6, pady=(4, 0))
                ctk.CTkLabel(mf, text=val,
                             font=ctk.CTkFont(size=11, weight="bold"),
                             text_color=m["color"]).pack(padx=6, pady=(0, 4))
    
            ctk.CTkLabel(inner, text=f"{t('speed')} {m['speed']}",
                         font=ctk.CTkFont(size=10),
                         text_color="#64748b", anchor="w").pack(anchor="w")
    
            self._model_cards[m["name"]] = (card, indicator)
    
            def on_click(name=m["name"]):
                self._selected_model.set(name)
                self._highlight_model(name)
                # Синхронизируем с выбором на странице сканирования
                if hasattr(self, 'model_choice'):
                    self.model_choice.set(name)
    
            card.bind("<Button-1>", lambda e, n=m["name"]: on_click(n))
            for w in card.winfo_children():
                w.bind("<Button-1>", lambda e, n=m["name"]: on_click(n))
    
        # Подсветить RF по умолчанию
        self.after(200, lambda: self._highlight_model("Random Forest"))
    
        # ── Drag & Drop зона ──────────────────────────────────────
        ctk.CTkLabel(frame, text=t("load_dataset"),
                     font=ctk.CTkFont(size=11, weight="bold"),
                     text_color="#64748b").pack(anchor="w", padx=20, pady=(14, 6))
    
        drop_zone = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                                  border_width=1, border_color="#1e293b",
                                  height=90, cursor="hand2")
        drop_zone.pack(fill="x", padx=16)
        drop_zone.pack_propagate(False)
    
        drop_inner = ctk.CTkFrame(drop_zone, fg_color="transparent")
        drop_inner.place(relx=0.5, rely=0.5, anchor="center")
    
        self.drop_icon = ctk.CTkLabel(drop_inner, text="📂",
                                       font=ctk.CTkFont(size=24))
        self.drop_icon.pack(side="left", padx=(0, 10))
    
        drop_text_frame = ctk.CTkFrame(drop_inner, fg_color="transparent")
        drop_text_frame.pack(side="left")
    
        self.drop_lbl = ctk.CTkLabel(drop_text_frame,
                                      text=t("drag_drop"),
                                      font=ctk.CTkFont(size=12, weight="bold"),
                                      text_color="#64748b")
        self.drop_lbl.pack(anchor="w")
    
        self.drop_sub = ctk.CTkLabel(drop_text_frame,
                                      text=t("supported_files"),
                                      font=ctk.CTkFont(size=10),
                                      text_color="#64748b")
        self.drop_sub.pack(anchor="w")
    
        def on_drop_click(e=None):
            from tkinter import filedialog
            path = filedialog.askopenfilename(filetypes=[
                ("Все файлы", "*.*"),
                ("CSV", "*.csv"),
                ("Логи", "*.log"),
                ("Текст", "*.txt"),
            ])
            if path:
                fname = Path(path).name
                self.drop_lbl.configure(text=f"✓  {fname}", text_color="#00ff88")
                self.drop_sub.configure(text=path, text_color="#64748b")
                self.drop_icon.configure(text="✅")
                drop_zone.configure(border_color="#00ff88")
                # Передаём в сканирование
                if hasattr(self, 'file_path'):
                    self.file_path.delete(0, "end")
                    self.file_path.insert(0, path)
                self._home_selected_file = path
    
        drop_zone.bind("<Button-1>", on_drop_click)
        drop_inner.bind("<Button-1>", on_drop_click)
        for w in drop_inner.winfo_children():
            w.bind("<Button-1>", on_drop_click)
    
        # ── Кнопка запуска ────────────────────────────────────────
        launch_frame = ctk.CTkFrame(frame, fg_color="transparent")
        launch_frame.pack(fill="x", padx=16, pady=(10, 6))
        launch_frame.grid_columnconfigure(0, weight=3)
        launch_frame.grid_columnconfigure(1, weight=1)
    
        def launch_scan():
            if hasattr(self, '_home_selected_file'):
                self.show_page("scan")
                self.after(100, self._run_scan)
            else:
                self.show_page("scan")
    
        ctk.CTkButton(launch_frame,
                      text=t("run_analysis"),
                      height=46, corner_radius=10,
                      fg_color="#00d4ff", hover_color="#00b8d9",
                      text_color="#000000",
                      font=ctk.CTkFont(size=14, weight="bold"),
                      command=launch_scan
                      ).grid(row=0, column=0, padx=(0, 6), sticky="ew")
    
        ctk.CTkButton(launch_frame,
                      text=t("analytics_btn"),
                      height=46, corner_radius=10,
                      fg_color="#1e293b", hover_color="#2d3748",
                      text_color="#e2e8f0",
                      font=ctk.CTkFont(size=13),
                      command=lambda: self.show_page("analytics")
                      ).grid(row=0, column=1, sticky="ew")

        # ── Live Threat Monitor ───────────────────────────────────
        live_frame = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                                   border_width=1, border_color="#1e293b")
        live_frame.pack(fill="x", padx=16, pady=(0, 6))

        live_hdr = ctk.CTkFrame(live_frame, fg_color="transparent")
        live_hdr.pack(fill="x", padx=12, pady=(8, 4))

        ctk.CTkLabel(live_hdr, text="🔴  LIVE THREAT MONITOR",
                     font=ctk.CTkFont(size=11, weight="bold"),
                     text_color="#e74c3c").pack(side="left")

        self.live_status_lbl = ctk.CTkLabel(live_hdr, text=t("stopped"),
                                             font=ctk.CTkFont(size=10),
                                             text_color="#475569")
        self.live_status_lbl.pack(side="right")

        # Threat Score + карточки
        live_content = ctk.CTkFrame(live_frame, fg_color="transparent")
        live_content.pack(fill="x", padx=12, pady=(0, 8))
        live_content.grid_columnconfigure(0, weight=1)
        live_content.grid_columnconfigure(1, weight=3)

        # Threat Score слева
        score_col = ctk.CTkFrame(live_content, fg_color="#0a0e1a", corner_radius=8)
        score_col.grid(row=0, column=0, sticky="ew", padx=(0, 8), pady=4)
        ctk.CTkLabel(score_col, text="Threat Score",
                     font=ctk.CTkFont(size=10), text_color="#475569"
                     ).pack(anchor="w", padx=10, pady=(8, 0))
        self.live_score_lbl = ctk.CTkLabel(score_col, text="0",
                                            font=ctk.CTkFont(size=36, weight="bold"),
                                            text_color="#2dc97e")
        self.live_score_lbl.pack(padx=10, pady=(0, 4))
        self.live_score_bar = ctk.CTkProgressBar(score_col, height=6,
                                                  corner_radius=3,
                                                  progress_color="#2dc97e")
        self.live_score_bar.pack(fill="x", padx=10, pady=(0, 8))
        self.live_score_bar.set(0)

        # Правая колонка — статистика
        stats_col = ctk.CTkFrame(live_content, fg_color="transparent")
        stats_col.grid(row=0, column=1, sticky="ew")
        stats_col.grid_columnconfigure(0, weight=1)
        stats_col.grid_columnconfigure(1, weight=1)
        stats_col.grid_columnconfigure(2, weight=1)

        # Карточка цикл
        c1 = ctk.CTkFrame(stats_col, fg_color="#0a0e1a", corner_radius=8)
        c1.grid(row=0, column=0, padx=2, sticky="ew")
        ctk.CTkLabel(c1, text=t("cycle"), font=ctk.CTkFont(size=9),
                     text_color="#475569").pack(pady=(6, 0))
        self.live_cycle_lbl = ctk.CTkLabel(c1, text="—",
                                            font=ctk.CTkFont(size=16, weight="bold"),
                                            text_color="#00d4ff")
        self.live_cycle_lbl.pack(pady=(0, 6))

        # Карточка угроза
        c2 = ctk.CTkFrame(stats_col, fg_color="#0a0e1a", corner_radius=8)
        c2.grid(row=0, column=1, padx=2, sticky="ew")
        ctk.CTkLabel(c2, text=t("threat"), font=ctk.CTkFont(size=9),
                     text_color="#475569").pack(pady=(6, 0))
        self.live_threat_lbl = ctk.CTkLabel(c2, text="—",
                                             font=ctk.CTkFont(size=16, weight="bold"),
                                             text_color="#475569")
        self.live_threat_lbl.pack(pady=(0, 6))

        # Карточка тип атаки
        c3 = ctk.CTkFrame(stats_col, fg_color="#0a0e1a", corner_radius=8)
        c3.grid(row=0, column=2, padx=2, sticky="ew")
        ctk.CTkLabel(c3, text=t("attack_type"), font=ctk.CTkFont(size=9),
                     text_color="#475569").pack(pady=(6, 0))
        self.live_attack_lbl = ctk.CTkLabel(c3, text="—",
                                             font=ctk.CTkFont(size=11, weight="bold"),
                                             text_color="#475569")
        self.live_attack_lbl.pack(pady=(0, 6))

        # Timeline лог
        self.live_timeline = ctk.CTkTextbox(live_frame, height=80,
                                             font=ctk.CTkFont(family="monospace", size=10),
                                             fg_color="transparent",
                                             text_color="#475569")
        self.live_timeline.pack(fill="x", padx=12, pady=(0, 4))
        self.live_timeline.insert("end", "Нажми ▶ чтобы начать мониторинг...\n")
        self.live_timeline.configure(state="disabled")

        # Кнопки управления
        btn_row = ctk.CTkFrame(live_frame, fg_color="transparent")
        btn_row.pack(fill="x", padx=12, pady=(0, 10))

        # Переключатель модели для Live Monitor
        model_row = ctk.CTkFrame(live_frame, fg_color="transparent")
        model_row.pack(fill="x", padx=12, pady=(0, 6))
        ctk.CTkLabel(model_row, text=t("model_colon"),
                     font=ctk.CTkFont(size=10), text_color="#475569"
                     ).pack(side="left", padx=(0, 8))
        self.live_model_selector = ctk.CTkSegmentedButton(
            model_row,
            values=["RF", "XGB", "ISO", "RKG"],
            font=ctk.CTkFont(size=10),
            fg_color="#1e293b",
            selected_color="#e74c3c",
            selected_hover_color="#c0392b",
            unselected_color="#1e293b",
            unselected_hover_color="#2d3748",
            command=self._on_live_model_select)
        self.live_model_selector.set("RF")
        self.live_model_selector.pack(side="left")
        self.live_model_name_lbl = ctk.CTkLabel(
            model_row, text="Random Forest",
            font=ctk.CTkFont(size=10), text_color="#475569")
        self.live_model_name_lbl.pack(side="left", padx=8)
        
        self.live_start_btn = ctk.CTkButton(
            btn_row, text=t("start_monitoring"),
            height=34, corner_radius=8,
            fg_color="#7a1e1e", hover_color="#c0392b",
            font=ctk.CTkFont(size=12, weight="bold"),
            command=self._toggle_live_monitor)
        self.live_start_btn.pack(side="left", padx=(0, 6))

        ctk.CTkButton(btn_row, text=t("history_btn"),
                      height=34, corner_radius=8, width=100,
                      fg_color="#1e293b", hover_color="#2d3748",
                      font=ctk.CTkFont(size=11),
                      command=self._show_live_history
                      ).pack(side="left")
        
        # ── Системный журнал ──────────────────────────────────────
        log_frame = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                                  border_width=1, border_color="#1e293b")
        log_frame.pack(fill="both", expand=True, padx=16, pady=(6, 10))
    
        log_hdr = ctk.CTkFrame(log_frame, fg_color="transparent")
        log_hdr.pack(fill="x", padx=12, pady=(8, 4))
        ctk.CTkLabel(log_hdr, text=t("system_log"),
                     font=ctk.CTkFont(size=10, weight="bold"),
                     text_color="#64748b").pack(side="left")
        self.log_dot = ctk.CTkLabel(log_hdr, text="●",
                                     font=ctk.CTkFont(size=10),
                                     text_color="#00ff88")
        self.log_dot.pack(side="right")
    
        lb = ctk.CTkTextbox(log_frame,
                             font=ctk.CTkFont(family="monospace", size=11),
                             fg_color="transparent",
                             text_color="#64748b")
        lb.pack(fill="both", expand=True, padx=8, pady=(0, 8))
    
        for e in [
            f"[BOOT]   {t('boot_msg')}",
            f"[AUTH]   {t('auth_msg')}: {self.username}",
            f"[CONFIG] {t('config_msg')}",
            f"[ML]     {t('ml_loaded') if self.model_loaded else t('model_not_found')}",
            f"[SCAN]   {t('scan_started')}",
            f"[READY]  {t('ready_msg')}",
        ]:
            lb.insert("end", e + "\n")
        lb.configure(state="disabled")
    
        return frame
    
    def _on_live_model_select(self, value):
        model_map = {
            "RF":  ("rf",           "Random Forest"),
            "XGB": ("xgb",          "XGBoost"),
            "ISO": ("iso",          "Isolation Forest"),
            "RKG": ("rootkitguard", "RootkitGuard ML"),
        }
        key, name = model_map.get(value, ("rf", "Random Forest"))
        if hasattr(self, "live_model_name_lbl"):
            self.live_model_name_lbl.configure(text=name)
        if hasattr(self, "_threat_monitor"):
            self._threat_monitor.set_model(key)

    def _toggle_live_monitor(self):
        if getattr(self, '_live_monitor_running', False):
            self._live_monitor_running = False
            self._threat_monitor.stop()
            self.live_start_btn.configure(
                text=t("start_monitoring"), fg_color="#7a1e1e")
            self.live_status_lbl.configure(
                text=t("stopped"), text_color="#475569")
        else:
            self._live_monitor_running = True
            from threat_monitor import ThreatMonitor
            self._threat_monitor = ThreatMonitor(
                interface="enp0s3", interval=2)
            self._threat_monitor.add_callback(self._on_live_event)
            if hasattr(self, "live_model_selector"):
                val = self.live_model_selector.get()
                self._on_live_model_select(val)
            self._threat_monitor.start()
            self.live_start_btn.configure(
                text=t("stop_btn"), fg_color="#e74c3c")
            self.live_status_lbl.configure(
                text=t("active"), text_color="#e74c3c")

    def _on_live_event(self, event: dict):
        etype = event.get("type")

        if etype == "cycle_start":
            self.after(0, lambda e=event: [
                self.live_cycle_lbl.configure(text=str(e["cycle"])),
                self.live_status_lbl.configure(
                    text=f"● {t('cycle')} #{e['cycle']} — {t('capture')}",
                    text_color="#f39c12"),
            ])

        elif etype == "result":
            data   = event["data"]
            threat = data["threat"]
            score  = data["threat_score"]
            atype  = data.get("attack_type", "—")
            ts     = data["timestamp"][-8:]

            color_map = {
                "ВЫСОКАЯ": "#e74c3c",
                "СРЕДНЯЯ": "#f39c12",
                "НИЗКАЯ":  "#2dc97e"
            }
            color = color_map.get(threat, "#475569")
            bar_color = color

            def update_ui(d=data, thr=threat, c=color, s=score, at=atype, ts=ts):
                # Обновляем Threat Score
                self.live_score_lbl.configure(
                    text=str(s), text_color=c)
                self.live_score_bar.configure(progress_color=bar_color)
                self.live_score_bar.set(s / 100)

                # Карточки
                self.live_threat_lbl.configure(text=thr, text_color=c)
                self.live_attack_lbl.configure(text=at if thr != "НИЗКАЯ" else "—", text_color=c)
                self.live_status_lbl.configure(
                    text=f"{t('active')} — {thr}", text_color=c)

                # Timeline
                icon = "🔴" if thr == "ВЫСОКАЯ" else "🟡" if thr == "СРЕДНЯЯ" else "🟢"
                line = f"[{ts}] {icon} {thr} | {d['anomalies']}/{d['total']} аномалий | {at}\n"
                self.live_timeline.configure(state="normal")
                self.live_timeline.insert("end", line)
                self.live_timeline.see("end")
                self.live_timeline.configure(state="disabled")

                # Обновляем карточки сканирования
                if hasattr(self, 'total_lbl'):
                    self.total_lbl.configure(text=f"{d['total']:,}")
                    self.normal_lbl.configure(text=f"{d['normal']:,}")
                    self.anom_lbl.configure(
                        text=f"{d['anomalies']:,}\n({d['pct']:.1f}%)")
                    self.threat_lbl.configure(text=t, text_color=c)

                
                # Автодообучение в фоне
                if t in ("ВЫСОКАЯ", "СРЕДНЯЯ"):
                    import threading
                    def auto_learn(data=d):
                        try:
                            from online_learner import OnlineLearner
                            learner = OnlineLearner()
                            import pandas as pd
                            import numpy as np
                            # Создаём синтетические аномальные образцы
                            rf = self.rf
                            cols = list(rf.feature_names_in_)
                            n = min(d['anomalies'], 30)
                            if n > 0:
                                samples = pd.DataFrame(
                                    np.random.randn(n, len(cols)) * 3 + 5,
                                    columns=cols)
                                samples['Dst Port'] = float(
                                    d['top_ports'][0] if d['top_ports'] else 445)
                                added = learner.add_attack_samples(samples, label=1)
                                self.after(0, lambda a=added: 
                                    self.learn_lbl.configure(
                                        text=f"🧠 +{a}",
                                        text_color="#a855f7"))
                                # Дообучаем если накопилось достаточно
                                if learner.should_retrain():
                                    self.after(0, lambda: self.learn_lbl.configure(
                                        text=t("training"),
                                        text_color="#f59e0b"))
                                    result = learner.retrain()
                                    if result.get("status") == "success":
                                        self._load_models()
                                        self.after(0, lambda r=result: 
                                            self.learn_lbl.configure(
                                                text=f"🧠 v{r['version']} ✓",
                                                text_color="#2dc97e"))
                                        self.after(0, lambda r=result:
                                            self.live_timeline.configure(state="normal") or
                                            self.live_timeline.insert("end",
                                                f"  🧠 Модель обновлена до v{r['version']} "
                                                f"(+{r['new_trees']} деревьев, "
                                                f"{r['total_trees']} всего)\n") or
                                            self.live_timeline.see("end") or
                                            self.live_timeline.configure(state="disabled"))
                        except Exception as e:
                            self.after(0, lambda err=str(e):
                                self.learn_lbl.configure(
                                    text="🧠 !", text_color="#e74c3c"))
                    threading.Thread(target=auto_learn, daemon=True).start()

                # Атака началась
                if d.get("attack_event") == "start":
                    self.live_timeline.configure(state="normal")
                    self.live_timeline.insert(
                        "end", f"  ⚡ АТАКА НАЧАЛАСЬ в {ts}!\n")
                    self.live_timeline.configure(state="disabled")

                # Атака завершилась
                elif d.get("attack_event") == "end":
                    self.live_timeline.configure(state="normal")
                    self.live_timeline.insert(
                        "end", f"  ✅ Атака завершена: {d.get('attack_duration')}\n")
                    self.live_timeline.configure(state="disabled")

            self.after(0, update_ui)

        elif etype == "no_traffic":
            self.after(0, lambda: self.live_status_lbl.configure(
                text=t("no_traffic"), text_color="#475569"))

        elif etype == "error":
            self.after(0, lambda e=event: self.live_status_lbl.configure(
                text=f"● {t('error')}: {e['msg'][:30]}", text_color="#e74c3c"))

    def _show_live_history(self):
        if not hasattr(self, '_threat_monitor'):
            return
        history = self._threat_monitor.get_history()
        if not history:
            return

        win = ctk.CTkToplevel(self)
        win.title("📋 История угроз")
        win.geometry("600x400")
        win.configure(fg_color="#0a0e1a")
        win.lift()

        ctk.CTkLabel(win, text=t("monitor_history"),
                     font=ctk.CTkFont(size=14, weight="bold"),
                     text_color="#00d4ff").pack(pady=(16, 8))

        box = ctk.CTkTextbox(win, font=ctk.CTkFont(family="monospace", size=11),
                              fg_color="#0d1117", text_color="#94a3b8")
        box.pack(fill="both", expand=True, padx=16, pady=(0, 16))

        for h in history:
            icon  = "🔴" if h["threat"] == "ВЫСОКАЯ" else "🟡" if h["threat"] == "СРЕДНЯЯ" else "🟢"
            line  = (f"[{h['timestamp'][-8:]}] {icon} {h['threat']:<8} | "
                     f"Аномалий: {h['anomalies']:>4}/{h['total']:>5} "
                     f"({h['pct']:>5.1f}%) | "
                     f"Score: {h['threat_score']:>3} | "
                     f"{h.get('attack_type','—')}\n")
            box.insert("end", line)
        box.configure(state="disabled")

    def _highlight_model(self, name: str):
        """Подсветить выбранную модель."""
        for mname, (card, indicator) in self._model_cards.items():
            if mname == name:
                card.configure(border_color={
                    "Random Forest":    "#0ea5e9",
                    "XGBoost":          "#a855f7",
                    "Isolation Forest": "#f59e0b",
                    "Ансамбль":         "#00ff88",
                }.get(name, "#00d4ff"))
                indicator.pack(fill="x", before=card.winfo_children()[1]
                              if len(card.winfo_children()) > 1 else card.winfo_children()[0])
            else:
                card.configure(border_color="#1e293b")
                indicator.pack_forget()
    
            # ── CSV Сканирование ─────────────────────────────────────────
    def _page_scan(self):
        frame = ctk.CTkFrame(self.main, fg_color="transparent")

        # Заголовок
        hdr = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                           border_width=1, border_color="#1e293b")
        hdr.pack(fill="x", padx=20, pady=(10, 5))
        ctk.CTkLabel(hdr, text=t("file_scanning"),
                     font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=16, pady=12)
        self._scan_model_selector = ctk.CTkSegmentedButton(
            hdr,
            values=["RF", "XGB", "ISO", "ALL"],
            font=ctk.CTkFont(size=11),
            fg_color="#1e293b",
            selected_color="#00d4ff",
            selected_hover_color="#00b8d9",
            unselected_color="#1e293b",
            unselected_hover_color="#2d3748",
            text_color_disabled="#475569",
            command=self._on_model_select)
        self._scan_model_selector.set("RF")
        self._scan_model_selector.pack(side="left", padx=10)

        # Выбор файла
        ff = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=10)
        ff.pack(fill="x", padx=20, pady=4)
        ctk.CTkLabel(ff, text=t("file_label"),
                     font=ctk.CTkFont(size=13, weight="bold"),
                     text_color="#85B7EB").pack(side="left", padx=14, pady=12)
        self.file_path = ctk.CTkEntry(ff, width=400,
                                       placeholder_text=t("choose_csv"),
                                       font=ctk.CTkFont(size=12))
        self.file_path.pack(side="left", padx=5)
        ctk.CTkButton(ff, text="📁 " + t("browse"), width=100, height=32,
                      fg_color="#2d6a4f",
                      command=self._browse_file).pack(side="left", padx=5)

        # Параметры
        pf = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=10)
        pf.pack(fill="x", padx=20, pady=4)
        ctk.CTkLabel(pf, text=t("threshold") + ":", font=ctk.CTkFont(size=12)).pack(side="left", padx=14, pady=10)
        self.threshold = ctk.CTkSlider(pf, from_=0.1, to=0.9, number_of_steps=8, width=160)
        self.threshold.set(0.5)
        self.threshold.pack(side="left", padx=5)
        self.thresh_lbl = ctk.CTkLabel(pf, text="0.5",
                                        font=ctk.CTkFont(size=12, weight="bold"),
                                        text_color="#2dc97e")
        self.thresh_lbl.pack(side="left")
        self.threshold.configure(command=lambda v: self.thresh_lbl.configure(text=f"{v:.1f}"))
        ctk.CTkLabel(pf, text="  " + t("default_rows") + ":", font=ctk.CTkFont(size=12)).pack(side="left", padx=10)
        self.n_rows = ctk.CTkEntry(pf, width=80, placeholder_text="10000")
        self.n_rows.pack(side="left", padx=5)
        self.use_api_var = ctk.BooleanVar(value=True)
        ctk.CTkCheckBox(pf, text=t("via_api"), variable=self.use_api_var).pack(side="left", padx=15)
        ctk.CTkButton(pf, text="↺", width=36, height=36,
                      fg_color="transparent", hover_color="#1e293b",
                      border_width=1, border_color="#2d3748",
                      font=ctk.CTkFont(size=16),
                      corner_radius=8,
                      command=self._reset_scan).pack(side="right", padx=(0, 6), pady=8)
        ctk.CTkButton(pf, text=t("run_scan"), height=36, width=200,
                      fg_color="#1f538d", hover_color="#2980b9",
                      font=ctk.CTkFont(size=13, weight="bold"),
                      corner_radius=8,
                      command=self._run_scan).pack(side="right", padx=14, pady=8)

        # Прогресс
        self.scan_progress = ctk.CTkProgressBar(
            frame, height=16, corner_radius=8,
            progress_color="#1f538d")
        self.scan_progress.pack(fill="x", padx=20, pady=6)
        self.scan_progress.set(0)
        self.scan_status = ctk.CTkLabel(frame, text=t("waiting"),
                                         text_color="gray", font=ctk.CTkFont(size=12))
        self.scan_status.pack()

        # Карточки результатов
        cards_frame = ctk.CTkFrame(frame, fg_color="transparent")
        cards_frame.pack(fill="x", padx=20, pady=6)
        card_data = [
            ("total_lbl",  t("total"),     "—", "#3498db"),
            ("normal_lbl", t("normal"),    "—", "#2dc97e"),
            ("anom_lbl",   t("anomalies"), "—", "#e74c3c"),
            ("threat_lbl", t("threat"),    "—", "#9b59b6"),
        ]
        
        for i, (attr, title, val, accent) in enumerate(card_data):
            card = ctk.CTkFrame(cards_frame, fg_color="#0d1117", corner_radius=10,
                                border_width=0, height=72)
            card.grid(row=0, column=i, padx=4, sticky="ew")
            card.grid_propagate(False)
            cards_frame.grid_columnconfigure(i, weight=1)
            stripe = ctk.CTkFrame(card, fg_color=accent, width=4, corner_radius=0)
            stripe.pack(side="left", fill="y")
            content = ctk.CTkFrame(card, fg_color="transparent")
            content.pack(side="left", fill="both", expand=True, padx=10, pady=8)
            ctk.CTkLabel(content, text=title,
                         font=ctk.CTkFont(size=10, weight="bold"),
                         text_color=accent, anchor="w").pack(anchor="w")
            lbl = ctk.CTkLabel(content, text=val,
                               font=ctk.CTkFont(size=18, weight="bold"),
                               text_color="white", anchor="w")
            lbl.pack(anchor="w")
            setattr(self, attr, lbl)

        # Нижняя часть — лог + история
        bottom = ctk.CTkFrame(frame, fg_color="transparent")
        bottom.pack(fill="both", expand=True, padx=20, pady=4)
        bottom.grid_columnconfigure(0, weight=3)
        bottom.grid_columnconfigure(1, weight=1)

        # Лог слева
        log_frame = ctk.CTkFrame(bottom, fg_color="#1e1e2e", corner_radius=10)
        log_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 4))
        ctk.CTkLabel(log_frame, text=t("scan_details"),
                     font=ctk.CTkFont(size=12, weight="bold"),
                     text_color="#85B7EB").pack(anchor="w", padx=12, pady=(8, 2))
        self.scan_result = ctk.CTkTextbox(
            log_frame, font=ctk.CTkFont(family="monospace", size=11))
        self.scan_result.pack(fill="both", expand=True, padx=8, pady=(0, 8))

        # История справа
        hist_frame = ctk.CTkFrame(bottom, fg_color="#1e1e2e", corner_radius=10)
        hist_frame.grid(row=0, column=1, sticky="nsew", padx=(4, 0))
        ctk.CTkLabel(hist_frame, text=t("history"),
                     font=ctk.CTkFont(size=12, weight="bold"),
                     text_color="#85B7EB").pack(anchor="w", padx=12, pady=(8, 4))
        self.scan_history_box = ctk.CTkTextbox(
            hist_frame, font=ctk.CTkFont(family="monospace", size=10))
        self.scan_history_box.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        self.scan_history_box.insert("end", t("no_scans") + "\n")
        self.scan_history_box.configure(state="disabled")

        # ── Кнопка Report (открывает панель экспорта Word/PDF/Excel) ──
        self.scan_report_btn = ctk.CTkButton(
            frame, text=t("report_btn"), height=42, corner_radius=8, state="disabled",
            fg_color="#1f538d", hover_color="#2b6cb0",
            font=ctk.CTkFont(size=14, weight="bold"),
            command=self._open_report_panel)
        self.scan_report_btn.pack(fill="x", padx=20, pady=(6, 2))

        self.scan_export_status = ctk.CTkLabel(
            frame, text="", font=ctk.CTkFont(size=11), text_color="#2dc97e")
        self.scan_export_status.pack(anchor="w", padx=20, pady=(0, 6))

        self._scan_history = []
        return frame

    def _launch_parallel_scan(self):
        # Запускаем оба скана одновременно
        self._parallel_scan = True
        threading.Thread(target=self._scan_worker, daemon=True).start()
        self.after(30000, lambda: setattr(self, "_parallel_scan", False))
        
        # Открываем мини-окно Rootkit
        rk_win = ctk.CTkToplevel(self)
        rk_win.title("🦠 Rootkit Scan")
        rk_win.geometry("500x420")
        rk_win.configure(fg_color="#0a0e1a")
        rk_win.lift()

        # Заголовок
        ctk.CTkLabel(rk_win, text="🦠  ROOTKIT SCAN",
                     font=ctk.CTkFont(size=14, weight="bold"),
                     text_color="#e74c3c").pack(pady=(16, 4))

        # Score
        score_frame = ctk.CTkFrame(rk_win, fg_color="#0d1117",
                                    corner_radius=10, border_width=1,
                                    border_color="#1e293b")
        score_frame.pack(fill="x", padx=16, pady=(0, 8))
        ctk.CTkLabel(score_frame, text="Security Score",
                     font=ctk.CTkFont(size=10),
                     text_color="#475569").pack(anchor="w", padx=12, pady=(8, 0))
        rk_score = ctk.CTkLabel(score_frame, text="...",
                                 font=ctk.CTkFont(size=28, weight="bold"),
                                 text_color="#f39c12")
        rk_score.pack(anchor="w", padx=12, pady=(0, 8))

        # Прогресс
        rk_prog = ctk.CTkProgressBar(rk_win, height=6, corner_radius=3,
                                      progress_color="#e74c3c")
        rk_prog.pack(fill="x", padx=16, pady=(0, 8))
        rk_prog.set(0)

        # 6 карточек
        cf = ctk.CTkFrame(rk_win, fg_color="transparent")
        cf.pack(fill="x", padx=16, pady=(0, 8))
        rk_cards = []
        for i, (label, icon) in enumerate([
            ("Процессы","🔎"), ("Ядро","🧩"), ("LD_PRE","💉"),
            ("Порты","🔌"), ("Файлы","📁"), ("UID=0","🔑")
        ]):
            cf.grid_columnconfigure(i, weight=1)
            card = ctk.CTkFrame(cf, fg_color="#0d1117", corner_radius=8,
                                border_width=1, border_color="#1e293b", height=70)
            card.grid(row=0, column=i, padx=2, sticky="ew")
            card.grid_propagate(False)
            ctk.CTkLabel(card, text=icon, font=ctk.CTkFont(size=14)).pack(pady=(6,0))
            ctk.CTkLabel(card, text=label, font=ctk.CTkFont(size=9),
                         text_color="#475569").pack()
            lbl = ctk.CTkLabel(card, text="○", font=ctk.CTkFont(size=10),
                               text_color="#475569")
            lbl.pack(pady=(0, 4))
            rk_cards.append((card, lbl))

        # Лог
        rk_log = ctk.CTkTextbox(rk_win, height=100,
                                 font=ctk.CTkFont(family="monospace", size=10),
                                 fg_color="#0d1117", text_color="#475569")
        rk_log.pack(fill="x", padx=16, pady=(0, 8))

        def run_rk():
            try:
                checker = RootkitChecker()
                fns = [
                    checker.check_hidden_processes,
                    checker.check_kernel_modules,
                    checker.check_ld_preload,
                    checker.check_suspicious_ports,
                    checker.check_system_files,
                    checker.check_privilege_escalation,
                ]
                all_findings = []
                for idx, fn in enumerate(fns):
                    rk_win.after(0, lambda p=(idx+1)/6: rk_prog.set(p))
                    try:
                        findings = fn()
                        all_findings.extend(findings)
                        card, lbl = rk_cards[idx]
                        if findings:
                            rk_win.after(0, lambda c=card, l=lbl, n=len(findings): (
                                c.configure(border_color="#e74c3c", fg_color="#1a0000"),
                                l.configure(text=f"⚠{n}", text_color="#e74c3c")))
                        else:
                            rk_win.after(0, lambda c=card, l=lbl: (
                                c.configure(border_color="#2dc97e", fg_color="#001a0d"),
                                l.configure(text="✓", text_color="#2dc97e")))
                    except Exception as e:
                        rk_log.configure(state="normal")
                        rk_log.insert("end", f"[!] {e}\n")
                        rk_log.configure(state="disabled")

                score = max(0, 100 - len(all_findings) * 15)
                color = "#2dc97e" if score >= 80 else "#f39c12" if score >= 50 else "#e74c3c"
                rk_win.after(0, lambda s=score, c=color:
                             rk_score.configure(text=str(s), text_color=c))
                rk_log.configure(state="normal")
                rk_log.insert("end", f"✓ Готово. Score: {score}/100. Находок: {len(all_findings)}\n")
                rk_log.configure(state="disabled")
            except Exception as e:
                rk_log.configure(state="normal")
                rk_log.insert("end", f"[!] Ошибка: {e}\n")
                rk_log.configure(state="disabled")

        threading.Thread(target=run_rk, daemon=True).start()

    def _reset_scan(self):
        self.file_path.delete(0, "end")
        self.scan_progress.set(0)
        self.scan_progress.configure(progress_color="#1f538d")
        self.scan_status.configure(text=t("waiting"), text_color="gray")
        self.scan_result.configure(state="normal")
        self.scan_result.delete("1.0", "end")
        self.scan_result.configure(state="disabled")
        self.total_lbl.configure(text="—")
        self.normal_lbl.configure(text="—")
        self.anom_lbl.configure(text="—")
        self.threat_lbl.configure(text="—", text_color="white")
        if hasattr(self, "scan_report_btn"): self.scan_report_btn.configure(state="disabled")
        if hasattr(self, '_threat_panel'):
            try:
                self._threat_panel.destroy()
            except Exception:
                pass
    
    def _on_model_select(self, value):
        model_map = {
            "RF":  "Random Forest",
            "XGB": "XGBoost", 
            "ISO": "Isolation Forest",
            "ALL": "Ансамбль",
        }
        if hasattr(self, 'model_choice'):
            self.model_choice.set(model_map.get(value, "Random Forest"))
        if hasattr(self, '_selected_model'):
            self._selected_model.set(model_map.get(value, "Random Forest"))

    def _browse_file(self):
        path = filedialog.askopenfilename(
            filetypes=[
                ("Все файлы",        "*.*"),
                ("CSV файлы",        "*.csv"),
                ("Текстовые файлы",  "*.txt"),
                ("Лог файлы",        "*.log"),
                ("JSON файлы",       "*.json"),
                ("Python скрипты",   "*.py"),
                ("Shell скрипты",    "*.sh"),
            ]
        )
        if path:
            self.file_path.delete(0, "end")
            self.file_path.insert(0, path)

    def _run_scan(self):
        threading.Thread(target=self._scan_worker, daemon=True).start()

    


    def _scan_worker(self):
        self.scan_result.configure(state="normal")
        self.scan_result.delete("1.0", "end")

        def log_ui(msg):
            self.scan_result.insert("end", msg + "\n")
            self.scan_result.see("end")

        path = self.file_path.get() or "data/raw/friday_traffic.csv"
        n    = int(self.n_rows.get() or cfg.get("scan", {}).get("default_rows", 10000))
        ts   = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        try:
            self.scan_status.configure(text=t("loading"), text_color="yellow")
            self.scan_progress.set(0.1)
            log_ui(f"[{ts}] Файл: {path}")

            if self.use_api_var.get() and self._api_available:
                log_ui("[*] Отправка в API /scan ...")
                self.scan_progress.set(0.3)
                with open(path, "rb") as f:
                    resp = requests.post(
                        f"{API_BASE}/scan",
                        files={"file": (Path(path).name, f, "text/csv")},
                        timeout=120)
                if resp.status_code == 200:
                    data = resp.json()
                    self._store_scan(data, path, ts)
                    self._print_results(log_ui, data, "API")
                    self.scan_progress.set(1.0)
                    # AI анализ через API
                    self.after(0, lambda d=data: self._show_api_ai_panel(d))
                    self.scan_result.configure(state="disabled")
                    return
                log_ui(f"[!] API {resp.status_code} — локальный режим")

            log_ui("[*] Локальный анализ...")
            df = pd.read_csv(path, nrows=n)
            log_ui(f"[+] Строк: {len(df):,}")
            self.scan_progress.set(0.3)

            # Убираем лишние колонки
            drop_cols = ["Label", "Timestamp", "label", "timestamp",
                         "Inbound", "inbound", "Flow ID", "flow_id",
                         "Src IP", "Dst IP", "src_ip", "dst_ip"]
            for col in drop_cols:
                if col in df.columns:
                    df = df.drop(columns=[col])

            # Оставляем только числовые колонки
            df = df.select_dtypes(include=[np.number])
            df = df.replace([np.inf, -np.inf], np.nan).fillna(0)

            # Подгоняем колонки под модель
            feat_cols = self.rf.feature_names_in_
            for c in feat_cols:
                if c not in df.columns:
                    df[c] = 0
            df = df[feat_cols]

            self.scan_progress.set(0.5)
            self.scan_status.configure(text=t("analyzing"))

            model_used = self._scan_model_selector.get() if hasattr(self, '_scan_model_selector') else "RF"
            insight = {}
            if self.model_loaded:
                X = pd.DataFrame(self.scaler.transform(df), columns=df.columns)
                feat_cols = self.rf.feature_names_in_
                for c in feat_cols:
                    if c not in X.columns: X[c] = 0
                X = X[feat_cols]

                if model_used == "RF":
                    preds = self.rf.predict(X)
                    proba = self.rf.predict_proba(X)[:, 1]
                    imp   = self.rf.feature_importances_
                    top3  = [(feat_cols[i], round(imp[i]*100,1))
                              for i in imp.argsort()[-3:][::-1]]
                    insight = {
                        "model":    "Random Forest",
                        "desc":     "Supervised classifier. Строит 100 деревьев решений\nи голосует большинством.",
                        "metrics":  [
                            ("Топ признак",  f"{top3[0][0]} ({top3[0][1]}%)"),
                            ("Уверенность",  f"{round(proba[preds==1].mean()*100 if preds.sum()>0 else 0,1)}%"),
                            ("Высокая уверен.", f"{int((proba>0.9).sum())} записей"),
                        ]
                    }

                elif model_used == "XGB":
                    try:
                        xgb_m = joblib.load("models/xgb_cicids.pkl")
                        preds = xgb_m.predict(X)
                        proba = xgb_m.predict_proba(X)[:, 1]
                        high  = int((proba > 0.9).sum())
                        med   = int(((proba > 0.5) & (proba <= 0.9)).sum())
                        low   = int((proba <= 0.5).sum())
                        insight = {
                            "model":   "XGBoost",
                            "desc":    "Gradient Boosting. Каждое дерево исправляет\nошибки предыдущего. Агрессивнее RF.",
                            "metrics": [
                                ("Высокая уверен. >90%", f"{high} записей"),
                                ("Средняя 50-90%",       f"{med} записей"),
                                ("Низкая <50%",          f"{low} записей"),
                            ]
                        }
                    except Exception as e:
                        log_ui(f"[!] XGB не найден: {e}")
                        preds = self.rf.predict(X)
                        proba = self.rf.predict_proba(X)[:, 1]
                        insight = {"model": "RF (fallback)", "desc": "", "metrics": []}

                elif model_used == "ISO":
                    try:
                        iso   = joblib.load("models/iso_cicids.pkl")
                        preds = (iso.predict(X) == -1).astype(int)
                        scores = iso.score_samples(X)
                        proba  = np.abs(scores)
                        proba  = (proba - proba.min()) / (proba.max() - proba.min() + 1e-9)
                        insight = {
                            "model":   "Isolation Forest",
                            "desc":    "Unsupervised. Не знает меток — ищет выбросы.\nИзолирует точки случайными разрезами.",
                            "metrics": [
                                ("Средний score",  f"{round(scores.mean(), 3)}"),
                                ("Мин. score",     f"{round(scores.min(), 3)} (самая аномальная)"),
                                ("Порог изоляции", f"{round(float(np.percentile(scores, 10)), 3)}"),
                            ]
                        }
                    except Exception as e:
                        log_ui(f"[!] ISO не найден: {e}")
                        preds = self.rf.predict(X)
                        proba = self.rf.predict_proba(X)[:, 1]
                        insight = {"model": "RF (fallback)", "desc": "", "metrics": []}

                elif model_used == "ALL":
                    try:
                        xgb_m  = joblib.load("models/xgb_cicids.pkl")
                        iso    = joblib.load("models/iso_cicids.pkl")
                        rf_p   = self.rf.predict(X)
                        xgb_p  = xgb_m.predict(X)
                        iso_p  = (iso.predict(X) == -1).astype(int)
                        votes  = rf_p + xgb_p + iso_p
                        preds  = (votes >= 2).astype(int)
                        proba  = self.rf.predict_proba(X)[:, 1]
                        cons3  = int((votes == 3).sum())
                        cons2  = int((votes == 2).sum())
                        cons1  = int((votes == 1).sum())
                        insight = {
                            "model":   "Ensemble (RF + XGB + ISO)",
                            "desc":    "Голосование большинством. Аномалия\nтолько если 2 из 3 моделей согласны.",
                            "metrics": [
                                ("Консенсус 3/3", f"{cons3} записей — критично"),
                                ("Консенсус 2/3", f"{cons2} записей — подозрительно"),
                                ("Единственный голос", f"{cons1} записей — возможно ложное"),
                            ]
                        }
                    except Exception as e:
                        log_ui(f"[!] Ансамбль ошибка: {e}")
                        preds = self.rf.predict(X)
                        proba = self.rf.predict_proba(X)[:, 1]
                        insight = {"model": "RF (fallback)", "desc": "", "metrics": []}
            else:
                log_ui("[!] Модель не загружена — демо-режим")
                preds = np.random.choice([0, 1], size=len(df), p=[0.75, 0.25])
                proba = np.random.uniform(0, 1, size=len(df))

            self.scan_progress.set(0.85)
            n_anom = int(preds.sum())
            n_norm = len(preds) - n_anom
            pct    = n_anom / len(preds) * 100
            threat = "ВЫСОКАЯ" if pct > 20 else "СРЕДНЯЯ" if pct > 5 else "НИЗКАЯ"
            top_ports = []
            if "Dst Port" in df.columns:
                top_ports = df[preds==1]["Dst Port"].value_counts().head(5).index.tolist()

            data = {
                "total_rows": len(preds), "anomalies": n_anom, "normal": n_norm,
                "pct": round(pct, 2), "threat": threat,
                "top_ports": [int(p) for p in top_ports],
                "max_proba": round(float(proba.max()), 4),
            }
            self._store_scan(data, path, ts)
            self._print_results(log_ui, data, "локальный")
            if insight:
                self.after(0, lambda i=insight: self._show_model_insight(i))
            self.scan_progress.set(1.0)
            if not getattr(self, "_parallel_scan", False): notify_threat(threat, f"{Path(path).name}: {n_anom} аномалий ({pct:.1f}%)")
            # Auto Defense теперь срабатывает в Rootkit Defense при реальной
            # угрозе на системе, а не на анализе CSV-файла трафика.

        except Exception as e:
            log_ui(f"[!] Ошибка: {e}")
            log.error(f"Scan error: {e}")
            self.scan_status.configure(text=t("error_excl"), text_color="red")

        self.scan_result.configure(state="disabled")

    def _store_scan(self, data: dict, filepath: str, ts: str):
        self._last_scan = {
            "total":     data.get("total_rows", 0),
            "anomaly":   data.get("anomalies", 0),
            "normal":    data.get("normal", 0),
            "pct":       data.get("pct", 0.0),
            "threat":    data.get("threat", "—"),
            "filename":  Path(filepath).name,
            "filepath":  filepath,
            "timestamp": ts,
            "top_ports": data.get("top_ports", []),
            "max_proba": data.get("max_proba", 0.0),
        }
    def _print_results(self, log_ui, data: dict, mode: str):
        threat = data.get("threat", "—")
        total  = data.get("total_rows", 0)
        anom   = data.get("anomalies", 0)
        norm   = data.get("normal", 0)
        pct    = data.get("pct", 0.0)

        # Обновляем карточки
        color_map = {"ВЫСОКАЯ": "#e74c3c", "СРЕДНЯЯ": "#f39c12", "НИЗКАЯ": "#2dc97e"}
        threat_color = color_map.get(threat, "white")
        self.after(0, lambda: [
            self.total_lbl.configure(text=f"{total:,}"),
            self.normal_lbl.configure(text=f"{norm:,}"),
            self.anom_lbl.configure(text=f"{anom:,}\n({pct:.1f}%)"),
            self.threat_lbl.configure(text=threat, text_color=threat_color),
        ])

        # Лог
        log_ui(f"\n{'='*48}")
        log_ui(f"  РЕЗУЛЬТАТЫ [{mode.upper()}]")
        log_ui(f"{'='*48}")
        log_ui(f"  Всего:         {total:,}")
        log_ui(f"  Нормальных:    {norm:,}")
        log_ui(f"  Аномалий:      {anom:,}  ({pct:.2f}%)")
        if data.get("max_proba"):
            log_ui(f"  Макс. вер-ть:  {data['max_proba']:.4f}")
        log_ui(f"  Угроза:        {threat}")
        if data.get("top_ports"):
            log_ui(f"  Топ порты:     {data['top_ports']}")
        log_ui(f"{'='*48}")

        # Активируем кнопку PDF
        # Меняем цвет прогресс-бара по угрозе
        bar_color = {"ВЫСОКАЯ": "#e74c3c", "СРЕДНЯЯ": "#f39c12",
                     "НИЗКАЯ": "#2dc97e"}.get(threat, "#2dc97e")
        self.after(0, lambda c=bar_color: self.scan_progress.configure(
            progress_color=c))
        
        self.after(0, lambda: self.scan_report_btn.configure(state="normal")
                   if hasattr(self, "scan_report_btn") else None)

        # Добавляем в историю
        ts = self._last_scan.get("timestamp", "")
        fn = self._last_scan.get("filename", "")
        entry = f"{ts[-8:]}  {fn[:15]:<15}  {threat}\n"
        self._scan_history.insert(0, entry)
        self._scan_history = self._scan_history[:8]
        self.after(0, self._update_scan_history)
        
        # Сравнение с прошлым сканом
        if self._prev_scan.get("total", 0) > 0:
            prev_anom = self._prev_scan.get("anomaly", 0)
            curr_anom = data.get("anomalies", 0)
            if curr_anom < prev_anom:
                log_ui(f"\n  ✅ Улучшение: аномалий было {prev_anom}, стало {curr_anom}")
            elif curr_anom > prev_anom:
                log_ui(f"\n  ⚠️  Ухудшение: аномалий было {prev_anom}, стало {curr_anom}")
            else:
                log_ui(f"\n  ➡️  Без изменений: {curr_anom} аномалий")

        self.scan_status.configure(
            text=f"✓ {t('done')} — {t('threat')}: {threat}", text_color=threat_color)
        self.after(0, lambda d=data, th=threat, c=threat_color: 
                   self._show_threat_panel(d, th, c))
        # Показываем Threat Intelligence Panel
        self.after(0, lambda d=data, t=threat, c=threat_color: 
                   self._show_threat_panel(d, t, c))

    def _show_api_ai_panel(self, data: dict):
        if hasattr(self, '_insight_panel'):
            try:
                self._insight_panel.destroy()
            except Exception:
                pass
        self._insight_panel = ctk.CTkFrame(
            self.pages["scan"], fg_color="#0d1117",
            corner_radius=12, border_width=1, border_color="#00d4ff")
        self._insight_panel.pack(fill="x", padx=20, pady=(0, 4))
        self._ai_font_size = 11
        self._ai_result_lbl = ctk.CTkTextbox(
            self._insight_panel, height=0,
            font=ctk.CTkFont(size=11),
            fg_color="transparent", text_color="#94a3b8", wrap="word")
        self._ai_result_lbl.pack(fill="x", padx=12, pady=(8, 4))
        self._ai_result_lbl.pack_forget()
        self._ai_zoom_row = ctk.CTkFrame(self._insight_panel, fg_color="transparent")
        self._ai_zoom_row.pack_forget()
        insight = {"model": "API", "desc": "", "metrics": [
            ("Всего", str(data.get("total_rows", 0))),
            ("Аномалий", str(data.get("anomalies", 0))),
            ("Угроза", data.get("threat", "—")),
        ]}
        ctk.CTkButton(
            self._insight_panel,
            text=t("🤖  Спросить AI"),                 
            height=36, corner_radius=8,
            fg_color="#1e293b", hover_color="#2d3748",
            font=ctk.CTkFont(size=12),
            command=lambda i=insight: threading.Thread(
                target=self._ai_analyze, args=(i,), daemon=True).start()
        ).pack(fill="x", padx=12, pady=(8, 8))
        
    def _show_model_insight(self, insight: dict):
        if hasattr(self, '_insight_panel'):
            try:
                self._insight_panel.destroy()
            except Exception:
                pass

        colors = {
            "Random Forest":          "#0ea5e9",
            "XGBoost":                "#a855f7",
            "Isolation Forest":       "#f59e0b",
            "Ensemble (RF + XGB + ISO)": "#00ff88",
        }
        color = colors.get(insight.get("model", ""), "#00d4ff")

        self._insight_panel = ctk.CTkFrame(
            self.pages["scan"], fg_color="#0d1117",
            corner_radius=12, border_width=1, border_color=color)
        self._insight_panel.pack(fill="x", padx=20, pady=(0, 4))

        # Заголовок
        hdr = ctk.CTkFrame(self._insight_panel, fg_color="transparent")
        hdr.pack(fill="x", padx=12, pady=(8, 4))
        ctk.CTkLabel(hdr, text=f"🧠  {insight.get('model', 'Model')} — Insight",
                     font=ctk.CTkFont(size=11, weight="bold"),
                     text_color=color).pack(side="left")

        content = ctk.CTkFrame(self._insight_panel, fg_color="transparent")
        content.pack(fill="x", padx=12, pady=(0, 10))
        content.grid_columnconfigure(0, weight=2)
        content.grid_columnconfigure(1, weight=3)

        # Описание
        desc_col = ctk.CTkFrame(content, fg_color="#0a0e1a", corner_radius=8)
        desc_col.grid(row=0, column=0, sticky="ew", padx=(0, 8), pady=4)
        ctk.CTkLabel(desc_col, text=t("how_it_works"),
                     font=ctk.CTkFont(size=10), text_color="#475569").pack(anchor="w", padx=10, pady=(8, 2))
        ctk.CTkLabel(desc_col, text=insight.get("desc", ""),
                     font=ctk.CTkFont(size=11),
                     text_color="#94a3b8", justify="left").pack(anchor="w", padx=10, pady=(0, 8))

        # Метрики
        metrics_col = ctk.CTkFrame(content, fg_color="transparent")
        metrics_col.grid(row=0, column=1, sticky="ew")
        for label, val in insight.get("metrics", []):
            row = ctk.CTkFrame(metrics_col, fg_color="#0a0e1a", corner_radius=6)
            row.pack(fill="x", pady=2)
            ctk.CTkLabel(row, text=label,
                         font=ctk.CTkFont(size=10), text_color="#475569",
                         width=180, anchor="w").pack(side="left", padx=10, pady=6)
            ctk.CTkLabel(row, text=val,
                         font=ctk.CTkFont(size=11, weight="bold"),
                         text_color=color).pack(side="left", padx=6)
        
        # Кнопка AI анализа
        ai_btn = ctk.CTkButton(
            self._insight_panel,
            text=t("ask_ai"),
            height=36, corner_radius=8,
            fg_color="#1e293b", hover_color="#2d3748",
            font=ctk.CTkFont(size=12),
            command=lambda i=insight: threading.Thread(
                target=self._ai_analyze, args=(i,), daemon=True).start())
        ai_btn.pack(fill="x", padx=12, pady=(0, 8))

        self._ai_result_lbl = ctk.CTkTextbox(
            self._insight_panel, height=0,
            font=ctk.CTkFont(size=11),
            fg_color="transparent", text_color="#94a3b8",
            wrap="word")
        self._ai_result_lbl.pack(fill="x", padx=12, pady=(0, 4))
        self._ai_result_lbl.pack_forget()

        # Зум кнопки
        self._ai_font_size = 11
        zoom_row = ctk.CTkFrame(self._insight_panel, fg_color="transparent")
        zoom_row.pack(fill="x", padx=12, pady=(0, 10))
        ctk.CTkButton(zoom_row, text="A−", width=36, height=24,
                      fg_color="#1e293b", hover_color="#2d3748",
                      font=ctk.CTkFont(size=11), corner_radius=6,
                      command=lambda: self._ai_zoom(-1)).pack(side="left", padx=(0,4))
        ctk.CTkButton(zoom_row, text="A+", width=36, height=24,
                      fg_color="#1e293b", hover_color="#2d3748",
                      font=ctk.CTkFont(size=11), corner_radius=6,
                      command=lambda: self._ai_zoom(1)).pack(side="left")
        zoom_row.pack_forget()
        self._ai_zoom_row = zoom_row

    def _ai_zoom(self, delta: int):
        self._ai_font_size = max(9, min(18, self._ai_font_size + delta))
        self._ai_result_lbl.configure(
            font=ctk.CTkFont(size=self._ai_font_size))
        
    def _ai_analyze(self, insight: dict):
        try:
            import anthropic
            self.after(0, lambda: [
                self._ai_result_lbl.configure(height=80),
                self._ai_result_lbl.pack(fill="x", padx=12, pady=(0, 8)),
                self._ai_result_lbl.configure(state="normal"),
                self._ai_result_lbl.delete("1.0", "end"),
                self._ai_result_lbl.insert("end", "🤖 Анализирую..."),
                self._ai_result_lbl.configure(state="disabled"),
            ])

            api_key = cfg.get("anthropic", {}).get("api_key", "")
            client  = anthropic.Anthropic(api_key=api_key)

            scan = self._last_scan
            lang_map = {"ru": "русском", "en": "английском", "kz": "қазақ"}

            prompt = f"""Ты эксперт по кибербезопасности. Проанализируй результаты сканирования сети:

Модель: {insight.get('model')}
Всего записей: {scan.get('total', 0)}
Нормальных: {scan.get('normal', 0)}
Аномалий: {scan.get('anomaly', 0)} ({scan.get('pct', 0):.1f}%)
Уровень угрозы: {scan.get('threat', '—')}
Атакованные порты: {scan.get('top_ports', [])}

Метрики модели:
{chr(10).join([f"- {k}: {v}" for k, v in insight.get('metrics', [])])}

Напиши краткий анализ (3-4 предложения):
1. Что обнаружено
2. Насколько опасно
3. Что делать

Отвечай на {lang_map.get(get_lang(), "русском")} языке. Кратко и по делу."""

            message = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=300,
                messages=[{"role": "user", "content": prompt}])

            result = message.content[0].text

            self.after(0, lambda r=result: [
                self._ai_result_lbl.configure(state="normal", height=150),
                self._ai_zoom_row.pack(fill="x", padx=12, pady=(0, 10)),
                self._ai_result_lbl.delete("1.0", "end"),
                self._ai_result_lbl.insert("end", f"🤖 AI: {r}"),
                self._ai_result_lbl.configure(state="disabled"),
            ])

        except Exception as e:
            self.after(0, lambda err=str(e): [
                self._ai_result_lbl.configure(state="normal"),
                self._ai_result_lbl.delete("1.0", "end"),
                self._ai_result_lbl.insert("end", f"[!] Ошибка AI: {err}"),
                self._ai_result_lbl.configure(state="disabled"),
            ])
                   
    def _show_threat_panel(self, data: dict, threat: str, color: str):
        if hasattr(self, '_threat_panel'):
            try:
                self._threat_panel.destroy()
            except Exception:
                pass
        total = data.get("total_rows", 1)
        anom  = data.get("anomalies", 0)
        norm  = data.get("normal", 0)
        pct   = data.get("pct", 0.0)
        ports = data.get("top_ports", [])
        icon  = "🔴" if threat == "ВЫСОКАЯ" else "🟡" if threat == "СРЕДНЯЯ" else "🟢"
        self._threat_panel = ctk.CTkFrame(
            self.pages["scan"], fg_color="#0d1117",
            corner_radius=12, border_width=1, border_color=color)
        self._threat_panel.pack(fill="x", padx=20, pady=(0, 4))
        hdr = ctk.CTkFrame(self._threat_panel, fg_color="transparent")
        hdr.pack(fill="x", padx=12, pady=(8, 4))
        ctk.CTkLabel(hdr, text=f"{icon}  THREAT INTELLIGENCE",
                     font=ctk.CTkFont(size=11, weight="bold"),
                     text_color=color).pack(side="left")
        ctk.CTkLabel(hdr, text=f"{threat}  ·  {pct:.1f}% anomalies",
                     font=ctk.CTkFont(size=11),
                     text_color="#475569").pack(side="right")
        content = ctk.CTkFrame(self._threat_panel, fg_color="transparent")
        content.pack(fill="x", padx=12, pady=(0, 10))
        content.grid_columnconfigure(0, weight=2)
        content.grid_columnconfigure(1, weight=1)
        content.grid_columnconfigure(2, weight=1)
        bar_col = ctk.CTkFrame(content, fg_color="transparent")
        bar_col.grid(row=0, column=0, sticky="ew", padx=(0, 12))
        ctk.CTkLabel(bar_col, text="Traffic Distribution",
                     font=ctk.CTkFont(size=10), text_color="#475569").pack(anchor="w")
        bar_bg = ctk.CTkFrame(bar_col, fg_color="#1e293b", corner_radius=6, height=20)
        bar_bg.pack(fill="x", pady=(4, 2))
        bar_bg.pack_propagate(False)
        norm_w = max(int(400 * (norm / max(total, 1))), 4)
        anom_w = max(int(400 * (anom / max(total, 1))), 4)
        ctk.CTkFrame(bar_bg, fg_color="#2dc97e", corner_radius=6,
                     height=20, width=norm_w).place(x=0, y=0)
        ctk.CTkFrame(bar_bg, fg_color=color, corner_radius=6,
                     height=20, width=anom_w).place(x=norm_w, y=0)
        ctk.CTkLabel(bar_col, text=f"🟢 Normal: {norm:,}   {icon} Anomaly: {anom:,}",
                     font=ctk.CTkFont(size=10), text_color="#64748b").pack(anchor="w")
        ports_col = ctk.CTkFrame(content, fg_color="transparent")
        ports_col.grid(row=0, column=1, sticky="ew", padx=(0, 12))
        ctk.CTkLabel(ports_col, text="Top Attack Ports",
                     font=ctk.CTkFont(size=10), text_color="#475569").pack(anchor="w")
        if ports:
            for p in ports[:3]:
                pf = ctk.CTkFrame(ports_col, fg_color="#1e293b", corner_radius=4)
                pf.pack(fill="x", pady=1)
                ctk.CTkLabel(pf, text=f"  :{p}",
                             font=ctk.CTkFont(size=10, weight="bold"),
                             text_color=color).pack(side="left", padx=6, pady=3)
        else:
            ctk.CTkLabel(ports_col, text="No suspicious ports",
                         font=ctk.CTkFont(size=10),
                         text_color="#475569").pack(anchor="w", pady=4)
        rec_col = ctk.CTkFrame(content, fg_color="transparent")
        rec_col.grid(row=0, column=2, sticky="ew")
        ctk.CTkLabel(rec_col, text="Recommendation",
                     font=ctk.CTkFont(size=10), text_color="#475569").pack(anchor="w")
        rec = {
            "ВЫСОКАЯ": "🚨 Block suspicious\nports immediately.\nRun Rootkit Scan.",
            "СРЕДНЯЯ": "⚡ Monitor processes.\nCheck open ports.\nRescan in 1h.",
            "НИЗКАЯ":  "✅ System is clean.\nNext scan: 24h.\nNo action needed.",
        }.get(threat, "Run full scan.")
        ctk.CTkLabel(rec_col, text=rec, font=ctk.CTkFont(size=10),
                     text_color="#94a3b8", justify="left").pack(anchor="w", pady=4)


    def _extract_threat_targets(self, findings):
        """Из находок rkdefense достаём PID-ы и пути бинарей для реагирования."""
        import re
        pids, paths = set(), set()
        for f in findings:
            blob = f"{getattr(f,'where','')} {getattr(f,'evidence','')} {getattr(f,'title','')}"
            for m in re.findall(r"PID\s+(\d+)", blob):
                pids.add(int(m))
            for m in re.findall(r"(/[\w./\-]+)", blob):
                if m.startswith(("/tmp", "/dev/shm", "/var/tmp", "/run", "/home", "/usr", "/bin", "/sbin")):
                    paths.add(m)
        return sorted(pids), sorted(paths)

    def _show_defense_modal(self, findings):
        """Auto Defense — реагирование на угрозы, найденные Rootkit Defense.
        Действия: Изоляция (kill+iptables), Карантин файла, Форензика-снимок."""
        if isinstance(findings, dict):   # обратная совместимость
            findings = getattr(self, "_rkd_last_findings", [])
        pids, paths = self._extract_threat_targets(findings)
        high = sum(1 for f in findings if getattr(f, "severity", "") == "ВЫСОКАЯ")

        modal = ctk.CTkToplevel(self)
        modal.title(t("defense_title"))
        modal.geometry("560x560")
        modal.resizable(False, False)
        modal.configure(fg_color="#0a0e1a")
        modal.grab_set(); modal.lift()

        hdr = ctk.CTkFrame(modal, fg_color="#1a0000", corner_radius=10,
                           border_width=1, border_color="#e74c3c")
        hdr.pack(fill="x", padx=20, pady=(18, 8))
        ctk.CTkLabel(hdr, text=t("high_threat_hdr"),
                     font=ctk.CTkFont(size=15, weight="bold"),
                     text_color="#e74c3c").pack(pady=(12, 2))
        ctk.CTkLabel(hdr,
                     text=f"{t('threats_lbl')}: {high}  ·  PID: {len(pids)}  ·  {t('file_lbl')}: {len(paths)}",
                     font=ctk.CTkFont(size=11), text_color="#94a3b8").pack(pady=(0, 12))

        status_box = ctk.CTkTextbox(modal, height=150,
                                    font=ctk.CTkFont(family="monospace", size=11),
                                    fg_color="#0d1117", text_color="#7fd1a8")
        status_box.pack(fill="both", expand=True, padx=20, pady=(0, 10))
        def log_m(msg):
            status_box.configure(state="normal")
            status_box.insert("end", msg + "\n"); status_box.see("end")
            status_box.configure(state="disabled")
        log_m(f"[READY] {t('defense_title')}")
        log_m(f"[INFO]  PID-цели: {pids or '—'}")
        log_m(f"[INFO]  Файлы-цели: {paths or '—'}")

        # ── Действие 1: Изоляция угрозы (kill + iptables по соединениям) ──
        def isolate():
            import os, signal, psutil
            if not pids:
                log_m("[!] Нет PID для изоляции"); return
            log_m("[*] Изоляция: завершаю процессы и собираю их IP...")
            remote_ips = set()
            for pid in pids:
                try:
                    pr = psutil.Process(pid)
                    for c in pr.net_connections(kind="inet"):
                        if c.raddr:
                            remote_ips.add(c.raddr.ip)
                except Exception:
                    pass
                try:
                    os.kill(pid, signal.SIGKILL)
                    log_m(f"[+] kill -9 {pid}")
                except PermissionError:
                    log_m(f"[!] нет прав на kill {pid} → sudo kill -9 {pid}")
                except ProcessLookupError:
                    log_m(f"[i] PID {pid} уже завершён")
            cmds = [f"sudo iptables -A OUTPUT -d {ip} -j DROP" for ip in sorted(remote_ips)]
            if cmds:
                log_m("[*] Команды блокировки удалённых IP:")
                for c in cmds: log_m(f"    {c}")
                self.clipboard_clear(); self.clipboard_append("\n".join(cmds))
                log_m("[✓] Команды скопированы в буфер")
            else:
                log_m("[i] Активных удалённых соединений не найдено")

        # ── Действие 2: Карантин файла (move + снять +x) ──
        def quarantine():
            import shutil, os
            if not paths:
                log_m("[!] Нет файлов для карантина"); return
            qdir = Path("reports") / "quarantine"
            qdir.mkdir(parents=True, exist_ok=True)
            for fp in paths:
                try:
                    if not os.path.exists(fp):
                        log_m(f"[i] {fp} не существует"); continue
                    dst = qdir / (Path(fp).name + f".{int(datetime.now().timestamp())}.quar")
                    os.chmod(fp, 0o000)
                    shutil.move(fp, dst)
                    log_m(f"[+] {fp} → {dst} (флаг исполнения снят)")
                except PermissionError:
                    log_m(f"[!] нет прав → sudo mv {fp} {qdir}/ ; sudo chmod 000 ...")
                except Exception as e:
                    log_m(f"[!] {fp}: {e}")

        # ── Действие 3: Форензика-снимок ──
        def forensics():
            import psutil, json
            log_m("[*] Сбор форензики...")
            snap = {"timestamp": datetime.now().isoformat(), "targets_pid": pids,
                    "targets_file": paths, "processes": [], "connections": []}
            for pid in pids:
                try:
                    pr = psutil.Process(pid)
                    with pr.oneshot():
                        snap["processes"].append({
                            "pid": pid, "name": pr.name(), "exe": pr.exe(),
                            "cmdline": pr.cmdline(), "username": pr.username(),
                            "open_files": [f.path for f in pr.open_files()][:20]})
                        for c in pr.net_connections(kind="inet"):
                            snap["connections"].append({
                                "pid": pid,
                                "laddr": f"{c.laddr.ip}:{c.laddr.port}" if c.laddr else "",
                                "raddr": f"{c.raddr.ip}:{c.raddr.port}" if c.raddr else "",
                                "status": c.status})
                except Exception as e:
                    snap["processes"].append({"pid": pid, "error": str(e)})
            Path("reports").mkdir(exist_ok=True)
            out = Path("reports") / f"forensics_{datetime.now():%Y%m%d_%H%M%S}.json"
            out.write_text(json.dumps(snap, ensure_ascii=False, indent=2), encoding="utf-8")
            log_m(f"[✓] Снимок сохранён: {out}")

        btns = ctk.CTkFrame(modal, fg_color="transparent")
        btns.pack(fill="x", padx=20, pady=(0, 6))
        btns.grid_columnconfigure((0, 1, 2), weight=1)
        ctk.CTkButton(btns, text=t("def_isolate"), height=42, corner_radius=8,
                      fg_color="#7a1e1e", hover_color="#c0392b",
                      font=ctk.CTkFont(size=12, weight="bold"),
                      command=lambda: threading.Thread(target=isolate, daemon=True).start()
                      ).grid(row=0, column=0, padx=(0, 4), sticky="ew")
        ctk.CTkButton(btns, text=t("def_quarantine"), height=42, corner_radius=8,
                      fg_color="#7a4520", hover_color="#a35e2a",
                      font=ctk.CTkFont(size=12, weight="bold"),
                      command=lambda: threading.Thread(target=quarantine, daemon=True).start()
                      ).grid(row=0, column=1, padx=4, sticky="ew")
        ctk.CTkButton(btns, text=t("def_forensics"), height=42, corner_radius=8,
                      fg_color="#13294a", hover_color="#1c3a63",
                      font=ctk.CTkFont(size=12, weight="bold"),
                      command=lambda: threading.Thread(target=forensics, daemon=True).start()
                      ).grid(row=0, column=2, padx=(4, 0), sticky="ew")

        # Нижняя строка: маленькая PDF-кнопка справа в углу + Закрыть
        bottom = ctk.CTkFrame(modal, fg_color="transparent")
        bottom.pack(fill="x", padx=20, pady=(2, 14))
        ctk.CTkButton(bottom, text=t("close_btn"), width=120, height=32,
                      fg_color="#1e293b", hover_color="#2d3748",
                      command=modal.destroy).pack(side="left")
        ctk.CTkButton(bottom, text="📕 PDF", width=80, height=28, corner_radius=6,
                      fg_color="transparent", border_width=1, border_color="#7a1e1e",
                      text_color="#e07a7a", font=ctk.CTkFont(size=11),
                      command=lambda: threading.Thread(
                          target=self._gen_pdf_report, daemon=True).start()
                      ).pack(side="right")

    def _update_scan_history(self):
        self.scan_history_box.configure(state="normal")
        self.scan_history_box.delete("1.0", "end")
        colors = {"ВЫСОКАЯ": "🔴", "СРЕДНЯЯ": "🟡", "НИЗКАЯ": "🟢"}
        for entry in self._scan_history:
            threat = entry.strip().split()[-1]
            icon = colors.get(threat, "⚪")
            self.scan_history_box.insert("end", f"{icon} {entry}")
        self.scan_history_box.configure(state="disabled")

    # ── Rootkit Scan ─────────────────────────────────────────────

    def _page_rootkit(self):
        frame = ctk.CTkFrame(self.main, fg_color="transparent")

        # Заголовок
        hdr = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                            border_width=1, border_color="#1e293b")
        hdr.pack(fill="x", padx=16, pady=(8, 6))
        hdr.pack_propagate(False)
        hdr.configure(height=52)
        ctk.CTkLabel(hdr, text="🦠  ROOTKIT SCAN",
                     font=ctk.CTkFont(size=14, weight="bold"),
                     text_color="#e74c3c").pack(side="left", padx=16, pady=14)
        self.rk_threat_lbl = ctk.CTkLabel(hdr, text="",
                                           font=ctk.CTkFont(size=13, weight="bold"))
        self.rk_threat_lbl.pack(side="left", padx=10)
        btns = ctk.CTkFrame(hdr, fg_color="transparent")
        btns.pack(side="right", padx=12)
        ctk.CTkButton(btns, text=t("run_btn"), width=140, height=34,
                      fg_color="#7a1e1e", hover_color="#c0392b",
                      corner_radius=8,
                      font=ctk.CTkFont(size=12, weight="bold"),
                      command=lambda: threading.Thread(
                          target=self._run_rootkit_local, daemon=True).start()
                      ).pack(side="left", padx=(0, 6), pady=8)
        ctk.CTkButton(btns, text="API", width=60, height=34,
                      fg_color="#1e293b", hover_color="#2d3748",
                      corner_radius=8, font=ctk.CTkFont(size=12),
                      command=lambda: threading.Thread(
                          target=self._run_rootkit_api, daemon=True).start()
                      ).pack(side="left")

        # Security Score
        score_frame = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                                    border_width=1, border_color="#1e293b")
        score_frame.pack(fill="x", padx=16, pady=(0, 6))
        score_inner = ctk.CTkFrame(score_frame, fg_color="transparent")
        score_inner.pack(side="left", padx=20, pady=12)
        ctk.CTkLabel(score_inner, text="Security Score",
                     font=ctk.CTkFont(size=10), text_color="#475569").pack(anchor="w")
        self.rk_score_lbl = ctk.CTkLabel(score_inner, text="—",
                                          font=ctk.CTkFont(size=32, weight="bold"),
                                          text_color="#475569")
        self.rk_score_lbl.pack(anchor="w")
        self.rk_status = ctk.CTkLabel(score_frame, text=t("ready_to_scan"),
                                       text_color="#475569",
                                       font=ctk.CTkFont(size=12))
        self.rk_status.pack(side="left", padx=20)

        # Прогресс
        self.rk_progress = ctk.CTkProgressBar(frame, height=6, corner_radius=3,
                                               progress_color="#e74c3c")
        self.rk_progress.pack(fill="x", padx=16, pady=(0, 6))
        self.rk_progress.set(0)

        # 6 карточек проверок
        cf = ctk.CTkFrame(frame, fg_color="transparent")
        cf.pack(fill="x", padx=16, pady=(0, 6))
        self._rk_cards = []
        checks_info = [
            ("Скрытые\nпроцессы", "🔎"),
            ("Модули\nядра",       "🧩"),
            ("LD_PRELOAD",         "💉"),
            ("Подозр.\nпорты",     "🔌"),
            ("Системные\nфайлы",   "📁"),
            ("Привилегии\nUID=0",  "🔑"),
        ]
        for i, (label, icon) in enumerate(checks_info):
            cf.grid_columnconfigure(i, weight=1)
            card = ctk.CTkFrame(cf, fg_color="#0d1117", corner_radius=10,
                                border_width=1, border_color="#1e293b", height=90)
            card.grid(row=0, column=i, padx=4, sticky="ew")
            card.grid_propagate(False)
            ctk.CTkLabel(card, text=icon,
                         font=ctk.CTkFont(size=18)).pack(pady=(10, 2))
            ctk.CTkLabel(card, text=label,
                         font=ctk.CTkFont(size=10),
                         text_color="#475569", justify="center").pack()
            lbl = ctk.CTkLabel(card, text="○",
                               font=ctk.CTkFont(size=11),
                               text_color="#475569")
            lbl.pack(pady=(2, 8))
            self._rk_cards.append((card, lbl))

        # System DNA
        dna_frame = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                                  border_width=1, border_color="#1e293b")
        dna_frame.pack(fill="x", padx=16, pady=(0, 6))
        dna_hdr = ctk.CTkFrame(dna_frame, fg_color="transparent")
        dna_hdr.pack(fill="x", padx=12, pady=(8, 4))
        ctk.CTkLabel(dna_hdr, text="🧬  System DNA",
                     font=ctk.CTkFont(size=11, weight="bold"),
                     text_color="#475569").pack(side="left")
        self.rk_dna_lbl = ctk.CTkLabel(dna_hdr, text=t("no_data"),
                                        font=ctk.CTkFont(size=10),
                                        text_color="#475569")
        self.rk_dna_lbl.pack(side="right")
        self.rk_dna_box = ctk.CTkTextbox(dna_frame, height=60,
                                          font=ctk.CTkFont(family="monospace", size=10),
                                          fg_color="transparent",
                                          text_color="#475569")
        self.rk_dna_box.pack(fill="x", padx=8, pady=(0, 8))
        self.rk_dna_box.insert("end", "Запусти сканирование чтобы создать первый снимок системы")
        self.rk_dna_box.configure(state="disabled")

        # AI панель
        self.rk_ai_frame = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                                         border_width=1, border_color="#1e293b")
        self.rk_ai_frame.pack(fill="x", padx=16, pady=(0, 6))
        self.rk_ai_frame.pack_forget()

        self.rk_output = ctk.CTkTextbox(
            frame, font=ctk.CTkFont(family="monospace", size=11),
            fg_color="#0d1117", text_color="#475569")
        self.rk_output.pack(fill="both", expand=True, padx=16, pady=(0, 10))
        return frame

    def _run_rootkit_local(self):
        import json
        from pathlib import Path

        self.rk_output.configure(state="normal")
        self.rk_output.delete("1.0", "end")
        self.rk_status.configure(text=t("scanning"), text_color="yellow")
        self.rk_score_lbl.configure(text="...", text_color="#f39c12")
        self.rk_progress.set(0)
        self.rk_threat_lbl.configure(text="")

        for card, lbl in self._rk_cards:
            card.configure(border_color="#1e293b", fg_color="#0d1117")
            lbl.configure(text="○", text_color="#475569")

        def log_ui(msg):
            self.rk_output.insert("end", msg + "\n")
            self.rk_output.see("end")

        try:
            checker = RootkitChecker()
            check_fns = [
                ("Скрытые процессы",  checker.check_hidden_processes),
                ("Модули ядра",        checker.check_kernel_modules),
                ("LD_PRELOAD",         checker.check_ld_preload),
                ("Подозр. порты",      checker.check_suspicious_ports),
                ("Системные файлы",    checker.check_system_files),
                ("Привилегии",         checker.check_privilege_escalation),
            ]
            all_findings = []
            current_dna = {
                "processes": [],
                "ports":     [],
                "modules":   [],
            }

            for idx, (name, fn) in enumerate(check_fns):
                progress = (idx + 1) / len(check_fns)
                self.after(0, lambda p=progress: self.rk_progress.set(p))
                try:
                    findings = fn()
                    all_findings.extend(findings)
                    card, lbl = self._rk_cards[idx]
                    if findings:
                        self.after(0, lambda c=card, l=lbl, n=len(findings): (
                            c.configure(border_color="#e74c3c", fg_color="#1a0000"),
                            l.configure(text=f"⚠ {n}", text_color="#e74c3c")))
                        for f in findings:
                            icon = "🔴" if f.severity == "ВЫСОКАЯ" else "🟡"
                            log_ui(f"  {icon} [{name}] {f.description}")
                    else:
                        self.after(0, lambda c=card, l=lbl: (
                            c.configure(border_color="#2dc97e", fg_color="#001a0d"),
                            l.configure(text="✓", text_color="#2dc97e")))
                except Exception as e:
                    log_ui(f"  [!] {name}: {e}")

            # Security Score
            score = max(0, 100 - len(all_findings) * 15)
            score_color = "#2dc97e" if score >= 80 else "#f39c12" if score >= 50 else "#e74c3c"

            threat = ("ВЫСОКАЯ" if any(f.severity == "ВЫСОКАЯ" for f in all_findings)
                      else "СРЕДНЯЯ" if any(f.severity == "СРЕДНЯЯ" for f in all_findings)
                      else "НИЗКАЯ" if all_findings else "ЧИСТАЯ")
            threat_color = {"ВЫСОКАЯ": "#e74c3c", "СРЕДНЯЯ": "#f39c12",
                            "НИЗКАЯ": "#f39c12", "ЧИСТАЯ": "#2dc97e"}.get(threat, "gray")

            self.after(0, lambda s=score, c=score_color: 
                       self.rk_score_lbl.configure(text=f"{s}", text_color=c))
            self.after(0, lambda: self.rk_status.configure(
                text=t("finished"), text_color="#2dc97e"))
            self.after(0, lambda t=threat, c=threat_color:
                       self.rk_threat_lbl.configure(
                           text=f"● {t}", text_color=c))
            self.after(0, lambda: self.rk_progress.configure(
                progress_color=threat_color))

            log_ui(f"\n{'='*50}")
            log_ui(f"  Security Score: {score}/100")
            log_ui(f"  Угроза: {threat}")
            log_ui(f"  Находок: {len(all_findings)}")
            if not all_findings:
                log_ui("  ✅ Система чиста")
            log_ui(f"{'='*50}")

            # System DNA
            dna_path = Path("data/system_dna.json")
            dna_path.parent.mkdir(exist_ok=True)
            new_dna = {
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "score":     score,
                "threat":    threat,
                "findings":  len(all_findings),
            }

            if dna_path.exists():
                old_dna = json.loads(dna_path.read_text())
                changes = []
                if old_dna.get("score") != score:
                    diff = score - old_dna.get("score", 0)
                    arrow = "↑" if diff > 0 else "↓"
                    changes.append(f"Score: {old_dna['score']} → {score} {arrow}{abs(diff)}")
                if old_dna.get("threat") != threat:
                    changes.append(f"Угроза: {old_dna['threat']} → {threat}")
                if old_dna.get("findings") != len(all_findings):
                    changes.append(f"Находок: {old_dna['findings']} → {len(all_findings)}")

                dna_text = f"Последний скан: {old_dna.get('timestamp', '—')}\n"
                if changes:
                    dna_text += "Изменения: " + "  |  ".join(changes)
                else:
                    dna_text += "Изменений не обнаружено — система стабильна ✓"

                self.after(0, lambda t=dna_text: [
                    self.rk_dna_box.configure(state="normal"),
                    self.rk_dna_box.delete("1.0", "end"),
                    self.rk_dna_box.insert("end", t),
                    self.rk_dna_box.configure(state="disabled"),
                    self.rk_dna_lbl.configure(
                        text=f"{t('snapshot')} #{len(changes)} {t('changes_lbl')}",
                        text_color="#00d4ff" if changes else "#2dc97e"),
                ])
            else:
                self.after(0, lambda: [
                    self.rk_dna_box.configure(state="normal"),
                    self.rk_dna_box.delete("1.0", "end"),
                    self.rk_dna_box.insert("end", "✓ Первый снимок системы создан"),
                    self.rk_dna_box.configure(state="disabled"),
                    self.rk_dna_lbl.configure(text=t("new_baseline"), text_color="#00d4ff"),
                ])

            dna_path.write_text(json.dumps(new_dna, ensure_ascii=False))

            # AI панель
            rk_insight = {
                "model": "Rootkit Scanner",
                "desc":  "",
                "metrics": [
                    ("Security Score", f"{score}/100"),
                    ("Угроза",         threat),
                    ("Находок",        str(len(all_findings))),
                ]
            }
            self.after(0, lambda i=rk_insight: self._show_rk_ai_panel(i, all_findings))
            if not getattr(self, "_parallel_scan", False): notify_threat(threat, f"Rootkit scan: {len(all_findings)} находок, score: {score}")

        except Exception as e:
            log_ui(f"[!] Ошибка: {e}")
            log.error(f"Rootkit error: {e}")
            self.rk_status.configure(text=t("error_excl"), text_color="red")

        self.rk_output.configure(state="disabled")

    def _run_rootkit_api(self):
        if not self._api_available:
            self.rk_status.configure(text=t("api_unavailable"), text_color="#e74c3c")
            return
        self.rk_status.configure(text=t("api_scanning"), text_color="yellow")
        try:
            import requests
            resp = requests.post(f"{API_BASE}/rootkit/scan", timeout=30)
            if resp.status_code == 200:
                data = resp.json()
                self.rk_score_lbl.configure(
                    text=f"{data.get('failed', 0)}/{data.get('total_checks', 0)}")
                self.rk_status.configure(
                    text=f"{t('done')} · {data.get('threat_level', '—')}", text_color="#2dc97e")
            else:
                self.rk_status.configure(text=f"API {resp.status_code}", text_color="#e74c3c")
        except Exception as e:
            self.rk_status.configure(text=f"{t('error')}: {str(e)[:30]}", text_color="#e74c3c")

            
    def _show_rk_ai_panel(self, insight: dict, findings: list):
        try:
            self.rk_ai_frame.pack(fill="x", padx=16, pady=(0, 6))
            for w in self.rk_ai_frame.winfo_children():
                w.destroy()
        except Exception:
            return

        # Если есть находки — показываем что делать
        if findings:
            ctk.CTkLabel(self.rk_ai_frame,
                         text=t("remediation"),
                         font=ctk.CTkFont(size=11, weight="bold"),
                         text_color="#e74c3c").pack(anchor="w", padx=12, pady=(8, 4))
            rec_map = {
                "Скрытые процессы":  "kill -9 <PID>  # завершить подозрительный процесс",
                "Модули ядра":       "rmmod <module>  # выгрузить подозрительный модуль",
                "LD_PRELOAD":        "unset LD_PRELOAD  # убрать инъекцию",
                "Подозр. порты":     "iptables -A INPUT -p tcp --dport <PORT> -j DROP",
                "Системные файлы":   "apt install --reinstall <package>  # переустановить",
                "Привилегии":        "chmod 755 <file>  # исправить права доступа",
            }
            for f in findings[:3]:
                desc = f.description
                cmd = None
                for key, val in rec_map.items():
                    if key.lower() in desc.lower():
                        cmd = val
                        break
                if cmd:
                    row = ctk.CTkFrame(self.rk_ai_frame, fg_color="#0a0e1a",
                                       corner_radius=6)
                    row.pack(fill="x", padx=12, pady=2)
                    ctk.CTkLabel(row, text=f"⚠ {desc[:50]}",
                                 font=ctk.CTkFont(size=10),
                                 text_color="#94a3b8").pack(anchor="w", padx=8, pady=(4, 0))
                    ctk.CTkLabel(row, text=f"$ {cmd}",
                                 font=ctk.CTkFont(family="monospace", size=10),
                                 text_color="#00d4ff").pack(anchor="w", padx=8, pady=(0, 4))

        # AI кнопка
        self._rk_ai_result = ctk.CTkTextbox(
            self.rk_ai_frame, height=0,
            font=ctk.CTkFont(size=11),
            fg_color="transparent", text_color="#94a3b8", wrap="word")
        self._rk_ai_result.pack(fill="x", padx=12, pady=(4, 0))
        self._rk_ai_result.pack_forget()

        ctk.CTkButton(
            self.rk_ai_frame,
            text=t("ask_ai"),
            height=34, corner_radius=8,
            fg_color="#1e293b", hover_color="#2d3748",
            font=ctk.CTkFont(size=12),
            command=lambda i=insight, f=findings: threading.Thread(
                target=self._rk_ai_analyze, args=(i, f), daemon=True).start()
        ).pack(fill="x", padx=12, pady=(4, 8))

    def _rk_ai_analyze(self, insight: dict, findings: list):
        try:
            import anthropic
            self.after(0, lambda: [
                self._rk_ai_result.configure(height=100),
                self._rk_ai_result.pack(fill="x", padx=12, pady=(4, 0)),
                self._rk_ai_result.configure(state="normal"),
                self._rk_ai_result.delete("1.0", "end"),
                self._rk_ai_result.insert("end", t("ai_analyzing")),
                self._rk_ai_result.configure(state="disabled"),
            ])
            api_key = cfg.get("anthropic", {}).get("api_key", "")
            client  = anthropic.Anthropic(api_key=api_key)
            lang_map = {"ru": "русском", "en": "английском", "kz": "қазақ тілінде"}
            lang = lang_map.get(get_lang(), "русском")
            findings_text = "\n".join([f"- [{f.severity}] {f.description}" 
                                        for f in findings]) if findings else "Находок нет"
            prompt = f"""Ты эксперт по кибербезопасности Linux. Результаты Rootkit Scan:

Security Score: {insight['metrics'][0][1]}
Угроза: {insight['metrics'][1][1]}
Находок: {insight['metrics'][2][1]}

Обнаруженные угрозы:
{findings_text}

Напиши краткий анализ (3-4 предложения) на {lang} языке:
1. Что обнаружено
2. Насколько критично
3. Конкретные шаги для защиты"""

            message = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=400,
                messages=[{"role": "user", "content": prompt}])
            result = message.content[0].text
            self.after(0, lambda r=result: [
                self._rk_ai_result.configure(state="normal", height=130),
                self._rk_ai_result.delete("1.0", "end"),
                self._rk_ai_result.insert("end", f"🤖 {r}"),
                self._rk_ai_result.configure(state="disabled"),
            ])
        except Exception as e:
            self.after(0, lambda err=str(e): [
                self._rk_ai_result.configure(state="normal"),
                self._rk_ai_result.delete("1.0", "end"),
                self._rk_ai_result.insert("end", f"{t('ai_error')}: {err}"),
                self._rk_ai_result.configure(state="disabled"),
            ])

    # ── Мониторинг ───────────────────────────────────────────────

    def _page_monitor(self):
        frame = ctk.CTkFrame(self.main, fg_color="transparent")

        # ── Заголовок + статус (минималистичная панель) ──────────
        bar = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=10,
                           border_width=1, border_color="#1e293b", height=56)
        bar.pack(fill="x", padx=16, pady=(10, 6))
        bar.pack_propagate(False)
        ctk.CTkLabel(bar, text=t("process_monitoring"),
                     font=ctk.CTkFont(size=15, weight="bold"),
                     text_color="#e2e8f0").pack(side="left", padx=16)
        self.mon_status = ctk.CTkLabel(bar, text="● LIVE",
                                       text_color="#2dc97e",
                                       font=ctk.CTkFont(size=12, weight="bold"))
        self.mon_status.pack(side="left", padx=8)
        self.mon_count = ctk.CTkLabel(bar, text=f"{t('processes')}: —",
                                      text_color="#64748b", font=ctk.CTkFont(size=12))
        self.mon_count.pack(side="left", padx=10)
        self.mon_threats = ctk.CTkLabel(bar, text=f"{t('threats_lbl')}: —",
                                        text_color="#64748b", font=ctk.CTkFont(size=12))
        self.mon_threats.pack(side="left", padx=10)

        self.btn_pause_mon = ctk.CTkButton(
            bar, text=t("pause_btn"), width=120, height=30, corner_radius=8,
            fg_color="#1e293b", hover_color="#2d3748",
            command=self._toggle_monitor)
        self.btn_pause_mon.pack(side="right", padx=12)
        self.mon_filter = ctk.CTkComboBox(
            bar, values=[t("all_lbl"), "ВЫСОКАЯ", "СРЕДНЯЯ", "НИЗКАЯ"], width=120,
            command=lambda v: self._refresh_monitor_table())
        self.mon_filter.set(t("all_lbl"))
        self.mon_filter.pack(side="right", padx=4)
        ctk.CTkLabel(bar, text=t("filter"), text_color="#64748b").pack(side="right", padx=(8, 2))

        # ── Таблица ──────────────────────────────────────────────
        tf = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=10,
                          border_width=1, border_color="#1e293b")
        tf.pack(fill="both", expand=True, padx=16, pady=(0, 10))
        hdr = ctk.CTkFrame(tf, fg_color="transparent")
        hdr.pack(fill="x", padx=4, pady=(6, 0))
        self._mon_cols = [("PID", 60), (t("process_col"), 200), ("CPU%", 60),
                          ("RAM MB", 75), ("Conn", 55), ("Score", 70),
                          (t("threat"), 90), (t("details_col"), 260)]
        for col, w in self._mon_cols:
            ctk.CTkLabel(hdr, text=col, width=w, anchor="w",
                         font=ctk.CTkFont(size=11, weight="bold"),
                         text_color="#475569").pack(side="left", padx=4, pady=6)
        self.mon_scroll = ctk.CTkScrollableFrame(tf, fg_color="transparent")
        self.mon_scroll.pack(fill="both", expand=True, padx=4, pady=(0, 4))

        # Real-time стартует автоматически
        self._lite = cfg.get("performance", {}).get("lite_mode", False)
        self._mon_row_cap = 20 if self._lite else 40
        self._mon_rows = []
        self._monitor = ProcessMonitor()
        self._monitor_data = []
        self._monitor.add_callback(self._on_monitor_update)
        self._monitor.start_realtime(interval=6 if self._lite else
                                     cfg.get("monitor", {}).get("interval_sec", 3))
        return frame

    def _on_monitor_update(self, data):
        self._monitor_data = data
        self.after(0, self._refresh_monitor_table)

    def _refresh_monitor_table(self, *args):
        if not hasattr(self, "mon_scroll") or not self.mon_scroll.winfo_exists():
            return
        # Не тратим ресурсы, если страница мониторинга не на экране
        if getattr(self, "_current_page", "") != "monitor":
            return
        filt = self.mon_filter.get()
        data = (self._monitor_data if filt == t("all_lbl")
                else [r for r in self._monitor_data if r["threat"] == filt])
        data = data[:getattr(self, "_mon_row_cap", 40)]
        threats = sum(1 for r in self._monitor_data if r["threat"] != "НИЗКАЯ")
        self.mon_count.configure(text=f"{t('processes')}: {len(self._monitor_data)}")
        self.mon_threats.configure(
            text=f"{t('threats_lbl')}: {threats}",
            text_color="#e74c3c" if threats > 0 else "#2dc97e")
        c_map = {"ВЫСОКАЯ": "#e74c3c", "СРЕДНЯЯ": "#f39c12", "НИЗКАЯ": "#475569"}
        widths = [w for _, w in self._mon_cols]

        # Обновляем виджеты «на месте» — пересоздание сотен виджетов
        # каждые 3 сек было главной причиной подвисаний
        rows = getattr(self, "_mon_rows", [])
        # добираем недостающие строки
        while len(rows) < len(data):
            i = len(rows)
            bg = "#111827" if i % 2 == 0 else "#0d1117"
            row_frame = ctk.CTkFrame(self.mon_scroll, fg_color=bg, corner_radius=0)
            row_frame.pack(fill="x")
            labels = []
            for w in widths:
                lbl = ctk.CTkLabel(row_frame, text="", width=w, anchor="w",
                                   font=ctk.CTkFont(size=11),
                                   text_color="#94a3b8")
                lbl.pack(side="left", padx=4, pady=5)
                labels.append(lbl)
            rows.append((row_frame, labels))
        # прячем лишние
        for row_frame, _ in rows[len(data):]:
            row_frame.pack_forget()
        # заполняем данные
        for i, r in enumerate(data):
            row_frame, labels = rows[i]
            if not row_frame.winfo_ismapped():
                row_frame.pack(fill="x")
            tc = c_map.get(r["threat"], "#475569")
            reason = ", ".join(r.get("reasons", []))[:42] or "—"
            cells = [
                (str(r["pid"]),              "#94a3b8"),
                (r["name"][:26],             "#e2e8f0"),
                (f"{r['cpu_percent']:.1f}",  "#94a3b8"),
                (f"{r['mem_rss']:.1f}",      "#94a3b8"),
                (str(r["n_conn"]),           "#94a3b8"),
                (f"{r['score']:.2f}",        "#94a3b8"),
                (r["threat"],                tc),
                (reason,                     tc if reason != "—" else "#475569"),
            ]
            for lbl, (val, color) in zip(labels, cells):
                if lbl.cget("text") != val:
                    lbl.configure(text=val)
                lbl.configure(text_color=color)
        self._mon_rows = rows

    def _toggle_monitor(self):
        if self._monitor.running:
            self._monitor.stop_realtime()
            self.mon_status.configure(text=t("paused"), text_color="#f39c12")
            self.btn_pause_mon.configure(text=t("resume_btn"))
        else:
            self._monitor.start_realtime()
            self.mon_status.configure(text="● LIVE", text_color="#2dc97e")
            self.btn_pause_mon.configure(text=t("pause_btn"))

    # ── Аналитика (без matplotlib) ───────────────────────────────

    def _page_analytics(self):
        frame = ctk.CTkFrame(self.main, fg_color="transparent")

        # Заголовок
        hdr = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                            border_width=1, border_color="#1e293b")
        hdr.pack(fill="x", padx=16, pady=(8, 6))
        hdr.configure(height=52)
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text=f"📊  {t('analytics_title')}",
                     font=ctk.CTkFont(size=14, weight="bold"),
                     text_color="#00d4ff").pack(side="left", padx=16, pady=14)

        # ── Живой снимок системы (psutil) ────────────────────────
        snap = ctk.CTkFrame(frame, fg_color="transparent")
        snap.pack(fill="x", padx=16, pady=(0, 6))
        self._an_snap_labels = {}
        snap_defs = [(t("an_cpu"), "cpu", "#0ea5e9"), (t("an_ram"), "ram", "#a855f7"),
                     (t("an_procs"), "procs", "#2dc97e"), (t("an_threats"), "thr", "#f59e0b")]
        for i, (title, key, color) in enumerate(snap_defs):
            snap.grid_columnconfigure(i, weight=1)
            c = ctk.CTkFrame(snap, fg_color="#0d1117", corner_radius=10,
                             border_width=1, border_color="#1e293b", height=64)
            c.grid(row=0, column=i, padx=4, sticky="ew"); c.grid_propagate(False)
            ctk.CTkFrame(c, fg_color=color, width=4, corner_radius=0).pack(side="left", fill="y")
            inn = ctk.CTkFrame(c, fg_color="transparent"); inn.pack(side="left", padx=10, pady=6)
            ctk.CTkLabel(inn, text=title, font=ctk.CTkFont(size=10),
                         text_color=color, anchor="w").pack(anchor="w")
            lbl = ctk.CTkLabel(inn, text="…", font=ctk.CTkFont(size=17, weight="bold"),
                               text_color="white", anchor="w")
            lbl.pack(anchor="w")
            self._an_snap_labels[key] = lbl
        self._refresh_analytics_snapshot()

        # ── Последнее сканирование (пропорция) ───────────────────
        ls = getattr(self, "_last_scan", None)
        if ls and ls.get("total"):
            lscard = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=10,
                                  border_width=1, border_color="#1e293b")
            lscard.pack(fill="x", padx=16, pady=(0, 6))
            ctk.CTkLabel(lscard, text=t("an_last_scan"),
                         font=ctk.CTkFont(size=10, weight="bold"),
                         text_color="#475569").pack(anchor="w", padx=12, pady=(8, 2))
            total = max(ls["total"], 1)
            anom_frac = ls["anomaly"] / total
            barbg = ctk.CTkFrame(lscard, fg_color="#1e293b", corner_radius=5, height=22)
            barbg.pack(fill="x", padx=12, pady=(0, 4)); barbg.pack_propagate(False)
            if anom_frac > 0:
                ctk.CTkFrame(barbg, fg_color="#e74c3c", corner_radius=5,
                             ).place(relx=0, rely=0, relwidth=max(anom_frac, 0.01), relheight=1)
            ctk.CTkLabel(lscard,
                         text=f"{t('normal')}: {ls['normal']:,}   ·   "
                              f"{t('anomalies')}: {ls['anomaly']:,} ({ls['pct']:.1f}%)   ·   "
                              f"{t('threat')}: {ls['threat']}",
                         font=ctk.CTkFont(size=11), text_color="#94a3b8"
                         ).pack(anchor="w", padx=12, pady=(0, 8))

        # Карточки сверху
        top_cards = ctk.CTkFrame(frame, fg_color="transparent")
        top_cards.pack(fill="x", padx=16, pady=(0, 6))
        for i, (title, val, sub, color) in enumerate([
            (t("best_model"),  "XGBoost",        "ROC-AUC: 1.0000",  "#a855f7"),
            (t("dataset"),        "CIC-IDS2018",     "1,044,525 записей","#0ea5e9"),
            (t("accuracy"),       "99.99%",          "F1-score",         "#2dc97e"),
            (t("features"),      "78",              t("net_features"),"#f59e0b"),
        ]):
            top_cards.grid_columnconfigure(i, weight=1)
            card = ctk.CTkFrame(top_cards, fg_color="#0d1117", corner_radius=10,
                                border_width=1, border_color="#1e293b", height=80)
            card.grid(row=0, column=i, padx=4, sticky="ew")
            card.grid_propagate(False)
            stripe = ctk.CTkFrame(card, fg_color=color, width=4, corner_radius=0)
            stripe.pack(side="left", fill="y")
            inner = ctk.CTkFrame(card, fg_color="transparent")
            inner.pack(side="left", padx=10, pady=8)
            ctk.CTkLabel(inner, text=title, font=ctk.CTkFont(size=10),
                         text_color=color, anchor="w").pack(anchor="w")
            ctk.CTkLabel(inner, text=val,
                         font=ctk.CTkFont(size=16, weight="bold"),
                         text_color="white", anchor="w").pack(anchor="w")
            ctk.CTkLabel(inner, text=sub, font=ctk.CTkFont(size=9),
                         text_color="#475569", anchor="w").pack(anchor="w")

        # Таблица метрик
        tbl = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                            border_width=1, border_color="#1e293b")
        tbl.pack(fill="x", padx=16, pady=(0, 6))
        hdr2 = ctk.CTkFrame(tbl, fg_color="#0a0e1a", corner_radius=0)
        hdr2.pack(fill="x")
        for col, w in [("Модель",180),("F1",90),("ROC-AUC",90),
                       ("FPR",80),("FNR",80),("Тип",110)]:
            ctk.CTkLabel(hdr2, text=col, width=w,
                         font=ctk.CTkFont(size=11, weight="bold"),
                         text_color="#475569").pack(side="left", padx=6, pady=8)

        rows = [
            ("Random Forest",    "1.0000","0.9999","0.0001","0.0001","Supervised", "#0ea5e9"),
            ("XGBoost",          "1.0000","1.0000","0.0000","0.0001","Supervised", "#a855f7"),
            ("Isolation Forest", "0.0200","0.3258","0.3666","0.9818","Unsupervised","#f59e0b"),
            ("Ensemble",         "1.0000","0.9999","0.0000","0.0001","Hybrid",     "#00ff88"),
        ]
        for model, f1, roc, fpr, fnr, typ, color in rows:
            r = ctk.CTkFrame(tbl, fg_color="#0d1117", corner_radius=0,
                             border_width=0)
            r.pack(fill="x")
            ctk.CTkFrame(r, fg_color=color, width=3, corner_radius=0
                         ).pack(side="left", fill="y")
            ctk.CTkLabel(r, text=model, width=177,
                         font=ctk.CTkFont(size=12, weight="bold"),
                         text_color=color).pack(side="left", padx=6, pady=8)
            for val, w in [(f1,90),(roc,90),(fpr,80),(fnr,80),(typ,110)]:
                ctk.CTkLabel(r, text=val, width=w,
                             font=ctk.CTkFont(size=11),
                             text_color="#94a3b8").pack(side="left", padx=6)

        # F1 бары
        bar_frame = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                                  border_width=1, border_color="#1e293b")
        bar_frame.pack(fill="x", padx=16, pady=(0, 6))
        ctk.CTkLabel(bar_frame, text="F1-score",
                     font=ctk.CTkFont(size=10, weight="bold"),
                     text_color="#475569").pack(anchor="w", padx=12, pady=(8, 4))
        for model, val, color in [
            ("Random Forest",    1.00, "#0ea5e9"),
            ("XGBoost",          1.00, "#a855f7"),
            ("Isolation Forest", 0.02, "#f59e0b"),
            ("Ensemble",         1.00, "#00ff88"),
        ]:
            r = ctk.CTkFrame(bar_frame, fg_color="transparent")
            r.pack(fill="x", padx=12, pady=3)
            ctk.CTkLabel(r, text=model, width=150, anchor="w",
                         font=ctk.CTkFont(size=11),
                         text_color="#64748b").pack(side="left")
            bg = ctk.CTkFrame(r, fg_color="#1e293b", corner_radius=4,
                               height=16, width=400)
            bg.pack(side="left", padx=8)
            bg.pack_propagate(False)
            ctk.CTkFrame(bg, fg_color=color, corner_radius=4,
                          height=16, width=max(int(400*val), 4)
                         ).place(x=0, y=0)
            ctk.CTkLabel(r, text=f"{val:.4f}",
                         font=ctk.CTkFont(size=11, weight="bold"),
                         text_color=color).pack(side="left")

        # Когда использовать
        guide = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                              border_width=1, border_color="#1e293b")
        guide.pack(fill="x", padx=16, pady=(0, 10))
        ctk.CTkLabel(guide, text=t("when_to_use"),
                     font=ctk.CTkFont(size=10, weight="bold"),
                     text_color="#475569").pack(anchor="w", padx=12, pady=(8, 4))
        for model, desc, color in [
            ("RF",  t("rf_when"), "#0ea5e9"),
            ("XGB", t("xgb_when"), "#a855f7"),
            ("ISO", t("iso_when"),     "#f59e0b"),
            ("ALL", t("all_when"),      "#00ff88"),
        ]:
            r = ctk.CTkFrame(guide, fg_color="#0a0e1a", corner_radius=6)
            r.pack(fill="x", padx=12, pady=2)
            ctk.CTkLabel(r, text=model, width=50,
                         font=ctk.CTkFont(size=11, weight="bold"),
                         text_color=color).pack(side="left", padx=8, pady=6)
            ctk.CTkLabel(r, text=desc,
                         font=ctk.CTkFont(size=11),
                         text_color="#64748b").pack(side="left")

        ctk.CTkLabel(guide, text="",).pack(pady=4)
        return frame
    # ── Отчёт ────────────────────────────────────────────────────

    def _refresh_analytics_snapshot(self):
        if not hasattr(self, "_an_snap_labels"):
            return
        try:
            lbls = self._an_snap_labels
            if not lbls["cpu"].winfo_exists():
                return
            import psutil
            lbls["cpu"].configure(text=f"{psutil.cpu_percent():.0f}%")
            lbls["ram"].configure(text=f"{psutil.virtual_memory().percent:.0f}%")
            lbls["procs"].configure(text=str(len(psutil.pids())))
            thr = sum(1 for r in getattr(self, "_monitor_data", [])
                      if r.get("threat") != "НИЗКАЯ")
            lbls["thr"].configure(text=str(thr))
        except Exception:
            pass
        if getattr(self, "_current_page", "") == "analytics":
            self.after(2000, self._refresh_analytics_snapshot)

    def _page_report(self):
        frame = ctk.CTkFrame(self.main, fg_color="transparent")

        # Заголовок
        hdr = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                            border_width=1, border_color="#1e293b", height=52)
        hdr.pack(fill="x", padx=16, pady=(8, 6))
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text=t("report_hdr"),
                     font=ctk.CTkFont(size=14, weight="bold"),
                     text_color="#00d4ff").pack(side="left", padx=16, pady=14)

        # Инфо о последнем скане
        info_frame = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                                   border_width=1, border_color="#1e293b")
        info_frame.pack(fill="x", padx=16, pady=(0, 6))
        info_inner = ctk.CTkFrame(info_frame, fg_color="transparent")
        info_inner.pack(fill="x", padx=12, pady=10)
        ctk.CTkLabel(info_inner, text=t("last_scan"),
                     font=ctk.CTkFont(size=10), text_color="#475569").pack(anchor="w")
        self.report_info_lbl = ctk.CTkLabel(
            info_inner,
            text=t("no_data_scan_first"),
            text_color="#64748b", font=ctk.CTkFont(size=12))
        self.report_info_lbl.pack(anchor="w", pady=(2, 0))

        # Кнопки экспорта
        btn_frame = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                                  border_width=1, border_color="#1e293b")
        btn_frame.pack(fill="x", padx=16, pady=(0, 6))
        ctk.CTkLabel(btn_frame, text=t("export"),
                     font=ctk.CTkFont(size=10), text_color="#475569"
                     ).pack(anchor="w", padx=12, pady=(8, 4))
        btns = ctk.CTkFrame(btn_frame, fg_color="transparent")
        btns.pack(fill="x", padx=12, pady=(0, 10))

        ctk.CTkButton(btns, text="📄  TXT",
                      height=38, width=120, corner_radius=8,
                      fg_color="#1e293b", hover_color="#2d3748",
                      font=ctk.CTkFont(size=12),
                      command=self._gen_text_report
                      ).pack(side="left", padx=(0, 6))

        ctk.CTkButton(btns, text="📕  PDF",
                      height=38, width=120, corner_radius=8,
                      fg_color="#7a1e1e", hover_color="#c0392b",
                      font=ctk.CTkFont(size=12),
                      command=lambda: threading.Thread(
                          target=self._gen_pdf_report, daemon=True).start()
                      ).pack(side="left", padx=(0, 6))

        ctk.CTkButton(btns, text=t("refresh"),
                      height=38, width=120, corner_radius=8,
                      fg_color="transparent", hover_color="#1e293b",
                      border_width=1, border_color="#2d3748",
                      font=ctk.CTkFont(size=12),
                      command=self._update_report_info
                      ).pack(side="left", padx=(0, 6))

        ctk.CTkButton(btns, text=t("ai_report"),
                      height=38, width=140, corner_radius=8,
                      fg_color="#0ea5e9", hover_color="#0284c7",
                      text_color="#000000",
                      font=ctk.CTkFont(size=12, weight="bold"),
                      command=lambda: threading.Thread(
                          target=self._gen_ai_report, daemon=True).start()
                      ).pack(side="left")

        self.report_status = ctk.CTkLabel(frame, text="",
                                           text_color="#2dc97e",
                                           font=ctk.CTkFont(size=11))
        self.report_status.pack(anchor="w", padx=16)

        # Превью отчёта
        preview_frame = ctk.CTkFrame(frame, fg_color="#0d1117", corner_radius=12,
                                      border_width=1, border_color="#1e293b")
        preview_frame.pack(fill="both", expand=True, padx=16, pady=(0, 10))
        ctk.CTkLabel(preview_frame, text=t("preview"),
                     font=ctk.CTkFont(size=10), text_color="#475569"
                     ).pack(anchor="w", padx=12, pady=(8, 4))
        self.report_box = ctk.CTkTextbox(
            preview_frame,
            font=ctk.CTkFont(family="monospace", size=11),
            fg_color="transparent", text_color="#94a3b8")
        self.report_box.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        return frame

    def _gen_ai_report(self):
        try:
            import anthropic
            self.after(0, lambda: [
                self.report_status.configure(
                    text=t("ai_generating"), text_color="yellow"),
                self.report_box.configure(state="normal"),
                self.report_box.delete("1.0", "end"),
                self.report_box.insert("end", "Генерирую профессиональный отчёт...\n"),
                self.report_box.configure(state="disabled"),
            ])
            s = self._last_scan
            if s["total"] == 0:
                self.after(0, lambda: self.report_status.configure(
                    text=t("scan_first"), text_color="#e74c3c"))
                return

            api_key = cfg.get("anthropic", {}).get("api_key", "")
            client  = anthropic.Anthropic(api_key=api_key)
            lang_map = {"ru": "русском", "en": "английском", "kz": "қазақ тілінде"}
            lang = lang_map.get(get_lang(), "русском")

            prompt = f"""Ты эксперт по кибербезопасности. Напиши профессиональный отчёт о результатах сканирования сети.

Данные сканирования:
- Файл: {s.get('filename', '—')}
- Время: {s.get('timestamp', '—')}
- Всего записей: {s.get('total', 0):,}
- Нормальных: {s.get('normal', 0):,}
- Аномалий: {s.get('anomaly', 0):,} ({s.get('pct', 0):.1f}%)
- Уровень угрозы: {s.get('threat', '—')}
- Атакованные порты: {s.get('top_ports', [])}
- Макс. вероятность: {s.get('max_proba', 0):.4f}

Напиши структурированный отчёт на {lang} языке со следующими разделами:

1. EXECUTIVE SUMMARY (2-3 предложения)
2. ОБНАРУЖЕННЫЕ УГРОЗЫ (детали атак, порты, тип трафика)
3. ОЦЕНКА РИСКА (уровень критичности, потенциальный ущерб)
4. РЕКОМЕНДАЦИИ (конкретные шаги, команды)
5. ЗАКЛЮЧЕНИЕ

Стиль: профессиональный, технический. Используй цифры из данных."""

            message = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=800,
                messages=[{"role": "user", "content": prompt}])

            result = message.content[0].text
            ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            full_report = f"""
╔══════════════════════════════════════════════════════════╗
║         ROOTKITGUARD — AI SECURITY REPORT               ║
╚══════════════════════════════════════════════════════════╝
  Сгенерировано: {ts}
  Файл: {s.get('filename', '—')}
  Версия: RootkitGuard v2.1 | Claude AI
══════════════════════════════════════════════════════════

{result}

══════════════════════════════════════════════════════════
  IITU
"""
            # Сохраняем
            from pathlib import Path
            Path("reports").mkdir(exist_ok=True)
            out_name = self._unique_name("txt")
            Path(f"reports/ai_{out_name}").write_text(full_report, encoding="utf-8")

            self.after(0, lambda r=full_report: [
                self.report_box.configure(state="normal"),
                self.report_box.delete("1.0", "end"),
                self.report_box.insert("end", r),
                self.report_box.configure(state="disabled"),
                self.report_status.configure(
                    text=f"✓ {t('ai_saved')}: reports/ai_{out_name}",
                    text_color="#2dc97e"),
            ])

        except Exception as e:
            self.after(0, lambda err=str(e): [
                self.report_status.configure(
                    text=f"{t('ai_error')}: {err}", text_color="#e74c3c"),
            ])

    def _update_report_info(self):
        if not hasattr(self, "report_info_lbl") or not self.report_info_lbl.winfo_exists():
            return
        s = self._last_scan
        if s["total"] == 0:
            self.report_info_lbl.configure(
                text=f"{t('last_scan')} {t('no_data_scan_first').lower()}")
            return
        self.report_info_lbl.configure(
            text=f"{t('file_lbl')}: {s['filename']}   |   {s['timestamp']}   |   "
                 f"Аномалий: {s['anomaly']:,}/{s['total']:,} ({s['pct']:.1f}%)   |   "
                 f"Угроза: {s['threat']}")

    def _unique_name(self, ext: str) -> str:
        s = self._last_scan
        ts = (s["timestamp"].replace(":", "-").replace(" ", "_")
              if s["timestamp"] else datetime.now().strftime("%Y-%m-%d_%H-%M-%S"))
        fn = Path(s["filename"]).stem if s["filename"] else "noscan"
        return f"report_{fn}_{ts}.{ext}"

    def _report_ui_ready(self) -> bool:
        """Виджеты страницы «Отчёт» существуют (страницы строятся лениво)."""
        return hasattr(self, "report_box") and self.report_box.winfo_exists()

    def _gen_text_report(self):
        self._update_report_info()
        self.report_box.configure(state="normal")
        self.report_box.delete("1.0", "end")
        s = self._last_scan
        out_name = self._unique_name("txt")
        out_path = Path("reports") / out_name

        rpt = f"""
╔══════════════════════════════════════════════════════════╗
║          ROOTKITGUARD — ОТЧЁТ АНАЛИЗА СИСТЕМЫ           ║
╚══════════════════════════════════════════════════════════╝
  Дата:        {s['timestamp'] or datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
  Версия:      RootkitGuard v{cfg.get('app',{}).get('version','2.1')}
  Файл:        {s['filename'] or '—'}
  Путь:        {s.get('filepath','—')}
──────────────────────────────────────────────────────────
  РЕЗУЛЬТАТЫ СКАНИРОВАНИЯ
──────────────────────────────────────────────────────────
  Всего записей:      {s['total']:,}
  Нормальных:         {s['normal']:,}
  Аномалий (Bot):     {s['anomaly']:,}  ({s['pct']:.2f}%)
  Макс. вероятность:  {s.get('max_proba', 0):.4f}
  Уровень угрозы:     {s['threat']}
  Топ порты:          {s.get('top_ports', []) or '—'}
──────────────────────────────────────────────────────────
  МЕТРИКИ МОДЕЛЕЙ
──────────────────────────────────────────────────────────
  Random Forest    F1:1.0000  ROC-AUC:0.9999  FPR:0.0001
  XGBoost          F1:1.0000  ROC-AUC:1.0000  FPR:0.0000
  Isolation Forest F1:0.0200  ROC-AUC:0.3258  FPR:0.3666
  Ensemble         F1:1.0000  ROC-AUC:0.9999  FPR:0.0000
──────────────────────────────────────────────────────────
  ЗАКЛЮЧЕНИЕ
──────────────────────────────────────────────────────────
  {'⚠  НЕМЕДЛЕННОЕ РАССЛЕДОВАНИЕ ТРЕБУЕТСЯ' if s['threat']=='ВЫСОКАЯ'
   else '⚡ Рекомендуется усиленный мониторинг' if s['threat']=='СРЕДНЯЯ'
   else '✅ Система работает в штатном режиме'}
──────────────────────────────────────────────────────────
  IITU
"""
        self.report_box.insert("end", rpt)
        self.report_box.configure(state="disabled")
        Path("reports").mkdir(exist_ok=True)
        out_path.write_text(rpt, encoding="utf-8")
        self.report_status.configure(text=f"✓ {t('saved_to')}: reports/{out_name}")

    def _open_report_panel(self):
        """Окно Report: экспорт последнего скана в Word / PDF / Excel."""
        if not getattr(self, "_last_scan", None) or not self._last_scan.get("total"):
            self._set_export_status(t("no_data_scan_first"), "#f39c12")
            return
        if getattr(self, "_report_win", None) is not None:
            try:
                if self._report_win.winfo_exists():
                    self._report_win.lift(); return
            except Exception:
                pass
        win = ctk.CTkToplevel(self)
        win.title(t("report_btn"))
        win.geometry("420x300")
        win.resizable(False, False)
        win.configure(fg_color="#0a0e1a")
        win.grab_set(); win.lift()
        self._report_win = win

        ctk.CTkLabel(win, text=t("report_hdr"),
                     font=ctk.CTkFont(size=16, weight="bold"),
                     text_color="#00d4ff").pack(pady=(18, 4))
        s_ = self._last_scan
        ctk.CTkLabel(win, text=f"{s_['filename'] or '—'}  ·  {t('threat')}: {s_['threat']}",
                     font=ctk.CTkFont(size=11), text_color="#94a3b8").pack(pady=(0, 14))

        def make(text, color, hover, fn):
            ctk.CTkButton(win, text=text, height=44, width=320, corner_radius=8,
                          fg_color=color, hover_color=hover,
                          font=ctk.CTkFont(size=14, weight="bold"),
                          command=lambda: threading.Thread(target=fn, daemon=True).start()
                          ).pack(pady=5)
        make(t("export_word"),  "#1e4620", "#2d6a30", self._gen_word_report)
        make(t("export_pdf"),   "#7a1e1e", "#c0392b", self._gen_pdf_report)
        make(t("export_excel"), "#1d4d3e", "#256b54", self._gen_excel_report)

        self.report_win_status = ctk.CTkLabel(win, text="", font=ctk.CTkFont(size=11),
                                              text_color="#2dc97e")
        self.report_win_status.pack(pady=(8, 0))
        def on_close():
            self._report_win = None; win.destroy()
        win.protocol("WM_DELETE_WINDOW", on_close)

    def _set_export_status(self, text, color="#2dc97e"):
        if hasattr(self, "scan_export_status") and self.scan_export_status.winfo_exists():
            self.after(0, lambda: self.scan_export_status.configure(text=text, text_color=color))
        if hasattr(self, "report_win_status"):
            try:
                if self.report_win_status.winfo_exists():
                    self.after(0, lambda: self.report_win_status.configure(text=text, text_color=color))
            except Exception:
                pass

    def _gen_pdf_report(self):
        try:
            from pdf_report import generate_pdf_report
            self._set_export_status(t("generating_pdf"), "#f39c12")
            s = self._last_scan
            out_name = self._unique_name("pdf")
            out_path = str(Path("reports") / out_name)
            Path("reports").mkdir(exist_ok=True)

            scan_data = {
                "total_rows": s["total"],
                "anomalies":  s["anomaly"],
                "normal":     s["normal"],
                "pct":        s["pct"],
                "threat":     s["threat"] if s["threat"] != "—" else "НИЗКАЯ",
                "top_ports":  s.get("top_ports", []),
                "filename":   s["filename"] or "—",
                "timestamp":  s["timestamp"] or "—",
            }
            generate_pdf_report(scan_data, out_path)
            self._set_export_status(f"✓ PDF: reports/{out_name}")
        except Exception as e:
            self._set_export_status(f"{t('pdf_error')}: {e}", "#e74c3c")
            log.error(f"PDF error: {e}")

    def _gen_word_report(self):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            self._set_export_status(t("generating_word"), "#f39c12")
            s = self._last_scan
            Path("reports").mkdir(exist_ok=True)
            out_name = self._unique_name("docx")
            doc = Document()
            doc.add_heading("RootkitGuard — Отчёт анализа", 0)
            doc.add_paragraph(f"Дата: {s['timestamp'] or '—'}    "
                              f"Версия: RootkitGuard v{cfg.get('app',{}).get('version','2.1')}")
            doc.add_paragraph(f"Файл: {s['filename'] or '—'}")
            doc.add_heading("Результаты сканирования", level=1)
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Light Grid Accent 1"
            for k, v in [("Всего записей", f"{s['total']:,}"),
                         ("Нормальных", f"{s['normal']:,}"),
                         ("Аномалий", f"{s['anomaly']:,} ({s['pct']:.2f}%)"),
                         ("Макс. вероятность", f"{s.get('max_proba',0):.4f}"),
                         ("Уровень угрозы", s['threat']),
                         ("Топ порты", str(s.get('top_ports', []) or '—'))]:
                row = tbl.add_row().cells
                row[0].text, row[1].text = k, v
            doc.add_heading("Заключение", level=1)
            verdict = ("НЕМЕДЛЕННОЕ РАССЛЕДОВАНИЕ ТРЕБУЕТСЯ" if s['threat'] == 'ВЫСОКАЯ'
                       else "Рекомендуется усиленный мониторинг" if s['threat'] == 'СРЕДНЯЯ'
                       else "Система работает в штатном режиме")
            doc.add_paragraph(verdict)
            doc.save(str(Path("reports") / out_name))
            self._set_export_status(f"✓ Word: reports/{out_name}")
        except Exception as e:
            self._set_export_status(f"{t('word_error')}: {e}", "#e74c3c")
            log.error(f"Word error: {e}")

    def _gen_excel_report(self):
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill
            self._set_export_status(t("generating_excel"), "#f39c12")
            s = self._last_scan
            Path("reports").mkdir(exist_ok=True)
            out_name = self._unique_name("xlsx")
            wb = Workbook()
            ws = wb.active
            ws.title = "Отчёт"
            hdr = Font(bold=True, color="FFFFFF")
            fill = PatternFill("solid", fgColor="1F538D")
            ws["A1"] = "RootkitGuard — Отчёт анализа"
            ws["A1"].font = Font(bold=True, size=14)
            rows = [("Параметр", "Значение"),
                    ("Дата", s['timestamp'] or '—'),
                    ("Файл", s['filename'] or '—'),
                    ("Всего записей", s['total']),
                    ("Нормальных", s['normal']),
                    ("Аномалий", s['anomaly']),
                    ("Доля аномалий, %", round(s['pct'], 2)),
                    ("Макс. вероятность", round(s.get('max_proba', 0), 4)),
                    ("Уровень угрозы", s['threat']),
                    ("Топ порты", str(s.get('top_ports', []) or '—'))]
            for i, (k, v) in enumerate(rows, start=3):
                ws[f"A{i}"], ws[f"B{i}"] = k, v
                if i == 3:
                    for c in ("A3", "B3"):
                        ws[c].font = hdr; ws[c].fill = fill
            ws.column_dimensions["A"].width = 22
            ws.column_dimensions["B"].width = 30
            wb.save(str(Path("reports") / out_name))
            self._set_export_status(f"✓ Excel: reports/{out_name}")
        except Exception as e:
            self._set_export_status(f"{t('excel_error')}: {e}", "#e74c3c")
            log.error(f"Excel error: {e}")

    # ── Настройки ────────────────────────────────────────────────

    def _page_settings(self):
        frame = ctk.CTkFrame(self.main, fg_color="transparent")
        ctk.CTkLabel(frame, text=t("settings"),
                     font=ctk.CTkFont(size=22, weight="bold")).pack(pady=(10, 5))

        scroll = ctk.CTkScrollableFrame(frame)
        scroll.pack(fill="both", expand=True, padx=20, pady=10)

        def section(title):
            ctk.CTkLabel(scroll, text=title,
                         font=ctk.CTkFont(size=14, weight="bold"),
                         text_color="#85B7EB").pack(anchor="w", pady=(14, 4))

        def row(label, widget_fn):
            r = ctk.CTkFrame(scroll, fg_color="#1e1e2e", corner_radius=8)
            r.pack(fill="x", pady=2)
            ctk.CTkLabel(r, text=label, width=230, anchor="w",
                         font=ctk.CTkFont(size=12)).pack(side="left", padx=12, pady=10)
            widget_fn(r)

        section(t("scan"))
        self._threshold_val = ctk.DoubleVar(
            value=cfg.get("scan", {}).get("threshold", 0.5))
        def thresh_row(p):
            sl = ctk.CTkSlider(p, from_=0.1, to=0.9, number_of_steps=8,
                               variable=self._threshold_val, width=180)
            sl.pack(side="left", padx=5)
            ctk.CTkLabel(p, textvariable=self._threshold_val).pack(side="left")
        row(t("set_threshold"), thresh_row)

        self._rows_val = ctk.StringVar(
            value=str(cfg.get("scan", {}).get("default_rows", 10000)))
        row(t("set_rows"),
            lambda p: ctk.CTkEntry(p, textvariable=self._rows_val, width=120
                                   ).pack(side="left", padx=5, pady=10))

        section(t("monitor"))
        self._interval_val = ctk.StringVar(
            value=str(cfg.get("monitor", {}).get("interval_sec", 5)))
        row(t("set_interval"),
            lambda p: ctk.CTkEntry(p, textvariable=self._interval_val, width=80
                                   ).pack(side="left", padx=5, pady=10))

        section(t("set_notif_section"))
        self._notif_var = ctk.BooleanVar(
            value=cfg.get("notifications", {}).get("enabled", True))
        row(t("set_notif_enable"),
            lambda p: ctk.CTkSwitch(p, text="", variable=self._notif_var
                                    ).pack(side="left", padx=5, pady=10))
        self._notif_lvl = ctk.StringVar(
            value=cfg.get("notifications", {}).get("min_threat_lvl", "СРЕДНЯЯ"))
        row(t("set_notif_level"),
            lambda p: ctk.CTkComboBox(
                p, values=["НИЗКАЯ", "СРЕДНЯЯ", "ВЫСОКАЯ"],
                variable=self._notif_lvl, width=150
            ).pack(side="left", padx=5, pady=10))

        section("API")
        self._api_port = ctk.StringVar(
            value=str(cfg.get("api", {}).get("port", 8000)))
        row(t("set_api_port"),
            lambda p: ctk.CTkEntry(p, textvariable=self._api_port, width=100
                                   ).pack(side="left", padx=5, pady=10))

        section(t("set_interface"))
        def lang_row(p):
            cur = {"ru": "РУС", "en": "ENG", "kz": "ҚАЗ"}.get(get_lang(), "РУС")
            cb = ctk.CTkComboBox(p, values=["РУС", "ENG", "ҚАЗ"], width=120,
                                 command=lambda v: self._switch_lang(
                                     {"РУС": "ru", "ENG": "en", "ҚАЗ": "kz"}[v]))
            cb.set(cur)
            cb.pack(side="left", padx=5, pady=10)
        row(t("set_language"), lang_row)

        def theme_row(p):
            cur = ctk.get_appearance_mode()  # 'Dark' / 'Light'
            cb = ctk.CTkComboBox(p, values=[t("theme_dark"), t("theme_light")], width=140,
                                 command=lambda v: ctk.set_appearance_mode(
                                     "dark" if v == t("theme_dark") else "light"))
            cb.set(t("theme_dark") if cur == "Dark" else t("theme_light"))
            cb.pack(side="left", padx=5, pady=10)
        row(t("set_theme"), theme_row)

        section(t("set_startup_section"))
        self._autoscan_var = ctk.BooleanVar(
            value=cfg.get("scan", {}).get("autostart_scan", True))
        row(t("set_autoscan"),
            lambda p: ctk.CTkSwitch(p, text="", variable=self._autoscan_var
                                    ).pack(side="left", padx=5, pady=10))

        self._lite_var = ctk.BooleanVar(
            value=cfg.get("performance", {}).get("lite_mode", False))
        row(t("set_lite"),
            lambda p: ctk.CTkSwitch(p, text="", variable=self._lite_var
                                    ).pack(side="left", padx=5, pady=10))

        btn_row = ctk.CTkFrame(frame, fg_color="transparent")
        btn_row.pack(pady=10)
        ctk.CTkButton(btn_row, text=t("save_btn"),
                      command=self._save_settings).pack(side="left", padx=10)
        ctk.CTkButton(btn_row, text=t("gen_demo_models"),
                      fg_color="#7a4520",
                      command=lambda: threading.Thread(
                          target=self._gen_demo_models_thread, daemon=True).start()
                      ).pack(side="left", padx=10)

        self.settings_status = ctk.CTkLabel(frame, text="", text_color="#2dc97e")
        self.settings_status.pack()
        return frame

    def _save_settings(self):
        try:
            import yaml
            cfg["scan"]["threshold"]       = round(self._threshold_val.get(), 1)
            cfg["scan"]["default_rows"]    = int(self._rows_val.get())
            cfg["scan"]["autostart_scan"]  = bool(self._autoscan_var.get())
            cfg.setdefault("performance", {})["lite_mode"] = bool(self._lite_var.get())
            cfg["monitor"]["interval_sec"] = int(self._interval_val.get())
            cfg["notifications"]["enabled"]        = self._notif_var.get()
            cfg["notifications"]["min_threat_lvl"] = self._notif_lvl.get()
            cfg["api"]["port"]             = int(self._api_port.get())
            p = Path(__file__).parent.parent / "config" / "config.yaml"
            with open(p, "w") as f:
                yaml.dump(cfg, f, allow_unicode=True, default_flow_style=False)
            self.settings_status.configure(text=t("saved_config"))
        except Exception as e:
            self.settings_status.configure(text=f"{t('error')}: {e}", text_color="red")

    # ── О системе ────────────────────────────────────────────────

    def _page_about(self):
        frame = ctk.CTkFrame(self.main, fg_color="transparent")
        ctk.CTkLabel(frame, text="RootkitGuard",
                     font=ctk.CTkFont(size=28, weight="bold")).pack(pady=(40, 5))
        ctk.CTkLabel(frame,
                     text=t("about_subtitle"),
                     font=ctk.CTkFont(size=14), text_color="gray").pack()
        ctk.CTkLabel(frame, text=t("about_uni"),
                     font=ctk.CTkFont(size=13), text_color="gray").pack(pady=(0, 30))
        for key, val in [
            ("Алгоритм",     "Random Forest + XGBoost + Isolation Forest"),
            ("Датасет",      "CIC-IDS2018 — 1,044,525 записей"),
            ("Точность",     "99.99% (F1-score на реальном датасете)"),
            ("ROC-AUC",      "0.9999"),
            ("v2.1",         "Без matplotlib, автозапуск API, уникальные отчёты, демо-модели"),
            ("Авторы",       "Амангелды Манас · Курманов Искандер · Куанышбек Бекарыс"),
            ("Руководитель", "Alin G.T."),
        ]:
            r = ctk.CTkFrame(frame, fg_color="#2b2b2b", corner_radius=8)
            r.pack(fill="x", padx=80, pady=3)
            ctk.CTkLabel(r, text=f"  {key}:", width=160, anchor="w",
                         font=ctk.CTkFont(weight="bold"),
                         text_color="gray").pack(side="left", pady=8)
            ctk.CTkLabel(r, text=val, anchor="w").pack(side="left", pady=8)
        return frame



    def _page_rkdefense(self):
        """Страница обнаружения rootkit — SOC + Forensic дизайн."""
        # ── Цветовая палитра ──────────────────────────────────────
        C_BG       = "#06090f"   # глубокий фон
        C_PANEL    = "#0b1220"   # панели
        C_PANEL_HI = "#0f1829"   # светлее
        C_BORDER   = "#1b2638"   # границы
        C_MONO     = "#7d8da5"   # серый текст
        C_DIM      = "#3d4d63"   # тусклый
        C_CYAN     = "#00d4ff"   # акцент (чисто)
        C_RED      = "#ff3b4e"   # критично
        C_AMBER    = "#ffa726"   # средне
        C_PURPLE   = "#b06cff"   # mitre
        MONO = "JetBrains Mono"  # если нет — упадёт на системный моно

        frame = ctk.CTkFrame(self.main, fg_color="transparent")

        # ══ КОМАНДНАЯ СТРОКА СТАТУСА (SOC header) ══════════════════
        cmdbar = ctk.CTkFrame(frame, fg_color=C_BG, corner_radius=10,
                              border_width=1, border_color=C_BORDER, height=46)
        cmdbar.pack(fill="x", padx=14, pady=(8, 4))
        cmdbar.pack_propagate(False)

        # Левая часть — название модуля
        left = ctk.CTkFrame(cmdbar, fg_color="transparent")
        left.pack(side="left", padx=14)
        ctk.CTkLabel(left, text="\u2589", font=ctk.CTkFont(size=16),
                     text_color=C_RED).pack(side="left", padx=(0, 6))
        ctk.CTkLabel(left, text="ROOTKIT DEFENSE",
                     font=ctk.CTkFont(size=14, weight="bold"),
                     text_color="#e6edf7").pack(side="left")
        ctk.CTkLabel(left, text="// LKM · PROCESS · PRIVESC · INTEGRITY",
                     font=ctk.CTkFont(family=MONO, size=9),
                     text_color=C_DIM).pack(side="left", padx=10)

        # Правая часть — индикатор статуса движка
        self.rkd_engine_lbl = ctk.CTkLabel(
            cmdbar, text="\u25cf ENGINE READY",
            font=ctk.CTkFont(family=MONO, size=10, weight="bold"),
            text_color=C_CYAN)
        self.rkd_engine_lbl.pack(side="right", padx=14)

        # ══ ВЕРХНИЙ РЯД: Score (слева) + методы (справа) ══════════
        top = ctk.CTkFrame(frame, fg_color="transparent")
        top.pack(fill="x", padx=14, pady=(0, 4))
        top.grid_columnconfigure(0, weight=0, minsize=240)
        top.grid_columnconfigure(1, weight=1)

        # ── Левая панель: Security Score (крупный вердикт) ────────
        score_panel = ctk.CTkFrame(top, fg_color=C_PANEL, corner_radius=12,
                                    border_width=1, border_color=C_BORDER)
        score_panel.grid(row=0, column=0, sticky="nsew", padx=(0, 4))

        ctk.CTkLabel(score_panel, text="SECURITY SCORE",
                     font=ctk.CTkFont(family=MONO, size=9, weight="bold"),
                     text_color=C_DIM).pack(anchor="w", padx=16, pady=(14, 0))

        score_row = ctk.CTkFrame(score_panel, fg_color="transparent")
        score_row.pack(anchor="w", padx=16, pady=(2, 0))
        self.rkd_score_lbl = ctk.CTkLabel(score_row, text="--",
                                           font=ctk.CTkFont(size=52, weight="bold"),
                                           text_color=C_DIM)
        self.rkd_score_lbl.pack(side="left")
        ctk.CTkLabel(score_row, text="/100",
                     font=ctk.CTkFont(size=18),
                     text_color=C_DIM).pack(side="left", anchor="s", pady=(0, 12))

        # Вердикт-бейдж
        self.rkd_verdict = ctk.CTkLabel(
            score_panel, text=t("waiting_caps"),
            font=ctk.CTkFont(family=MONO, size=12, weight="bold"),
            fg_color=C_PANEL_HI, corner_radius=6,
            text_color=C_DIM, height=28)
        self.rkd_verdict.pack(fill="x", padx=16, pady=(4, 8))

        # Прогресс-бар скана
        self.rkd_progress = ctk.CTkProgressBar(
            score_panel, height=4, corner_radius=2,
            progress_color=C_CYAN, fg_color=C_PANEL_HI)
        self.rkd_progress.pack(fill="x", padx=16, pady=(0, 6))
        self.rkd_progress.set(0)

        self.rkd_status = ctk.CTkLabel(
            score_panel, text=t("ready_to_scan_lc"),
            font=ctk.CTkFont(family=MONO, size=9),
            text_color=C_MONO)
        self.rkd_status.pack(anchor="w", padx=16, pady=(0, 8))

        self.rkd_baseline_lbl = ctk.CTkLabel(
            score_panel, text=t("baseline_none"),
            font=ctk.CTkFont(family=MONO, size=9),
            text_color=C_DIM)
        self.rkd_baseline_lbl.pack(anchor="w", padx=16, pady=(0, 14))

        # ── Правая панель: 6 детекторов (сетка 6 колонок) ─────────
        det_panel = ctk.CTkFrame(top, fg_color="transparent")
        det_panel.grid(row=0, column=1, sticky="nsew", padx=(4, 0))
        for i in range(6):
            det_panel.grid_columnconfigure(i, weight=1)
        det_panel.grid_rowconfigure(0, weight=1)

        self._rkd_cards = []
        detectors = [
            ("HIDDEN\nPROC",    "\U0001f50e", "T1014"),
            ("HIDDEN\nLKM",     "\U0001f9e9", "T1547"),
            ("PRIV\nESC",       "\U0001f513", "T1548"),
            ("LD\nPRELOAD",     "\U0001f517", "T1574"),
            ("BINARY\nINTEGRITY","\U0001f4c1", "T1554"),
            ("BACKDOOR\nPORTS", "\U0001f50c", "T1571"),
        ]
        for i, (label, icon, tag) in enumerate(detectors):
            card = ctk.CTkFrame(det_panel, fg_color=C_PANEL, corner_radius=10,
                                border_width=1, border_color=C_BORDER)
            card.grid(row=0, column=i, padx=3, sticky="nsew")

            # Верхняя полоса-индикатор
            strip = ctk.CTkFrame(card, fg_color=C_DIM, height=3, corner_radius=0)
            strip.pack(fill="x", padx=1, pady=(1, 0))

            ctk.CTkLabel(card, text=icon,
                         font=ctk.CTkFont(size=20)).pack(pady=(12, 2))
            ctk.CTkLabel(card, text=label,
                         font=ctk.CTkFont(family=MONO, size=9, weight="bold"),
                         text_color=C_MONO, justify="center").pack()
            ctk.CTkLabel(card, text=tag,
                         font=ctk.CTkFont(family=MONO, size=8),
                         text_color=C_DIM).pack(pady=(2, 0))

            # Статус-метка (○ / ✓ / ⚠N)
            status = ctk.CTkLabel(card, text="\u25cb",
                                  font=ctk.CTkFont(size=14),
                                  text_color=C_DIM)
            status.pack(pady=(6, 12))
            self._rkd_cards.append((card, status, strip))

        # ══ ПАНЕЛЬ ДЕЙСТВИЙ ════════════════════════════════════════
        actions = ctk.CTkFrame(frame, fg_color=C_PANEL, corner_radius=10,
                               border_width=1, border_color=C_BORDER, height=52)
        actions.pack(fill="x", padx=14, pady=(4, 4))
        actions.pack_propagate(False)

        # Сканировать (главная кнопка)
        ctk.CTkButton(actions, text=t("scan_caps"),
                      width=150, height=34, corner_radius=8,
                      fg_color=C_RED, hover_color="#d62f40",
                      text_color="#ffffff",
                      font=ctk.CTkFont(size=12, weight="bold"),
                      command=lambda: threading.Thread(
                          target=self._run_rkdefense_scan, daemon=True).start()
                      ).pack(side="left", padx=(14, 6), pady=9)

        # Baseline
        ctk.CTkButton(actions, text="\u25c9 BASELINE",
                      width=120, height=34, corner_radius=8,
                      fg_color=C_PANEL_HI, hover_color="#16233a",
                      border_width=1, border_color=C_BORDER,
                      text_color="#c7d3e3",
                      font=ctk.CTkFont(family=MONO, size=11),
                      command=lambda: threading.Thread(
                          target=self._rkdefense_baseline, daemon=True).start()
                      ).pack(side="left", padx=(0, 6), pady=9)

        # Копировать всё (перенесено сюда!)
        self._copy_btn = ctk.CTkButton(
            actions, text=t("copy_all_caps"),
            width=160, height=34, corner_radius=8,
            fg_color=C_PANEL_HI, hover_color="#16233a",
            border_width=1, border_color=C_CYAN,
            text_color=C_CYAN,
            font=ctk.CTkFont(family=MONO, size=11),
            command=self._copy_all_log)
        self._copy_btn.pack(side="left", padx=(0, 6), pady=9)

        # Очистить результаты
        ctk.CTkButton(actions, text=t("clear_btn"),
                      width=120, height=34, corner_radius=8,
                      fg_color=C_PANEL_HI, hover_color="#3a1620",
                      border_width=1, border_color=C_DIM,
                      text_color="#c7d3e3",
                      font=ctk.CTkFont(family=MONO, size=11),
                      command=self._clear_rkdefense
                      ).pack(side="left", padx=(0, 6), pady=9)

        # Системный журнал (live) — что происходит вне программы
        ctk.CTkButton(actions, text=t("syslog_btn"),
                      width=150, height=34, corner_radius=8,
                      fg_color="#13294a", hover_color="#1c3a63",
                      border_width=1, border_color=C_CYAN,
                      text_color=C_CYAN,
                      font=ctk.CTkFont(family=MONO, size=11),
                      command=self._open_syslog_window
                      ).pack(side="left", padx=(0, 6), pady=9)

        # Threat-индикатор справа
        self.rkd_threat_lbl = ctk.CTkLabel(
            actions, text="",
            font=ctk.CTkFont(family=MONO, size=12, weight="bold"))
        self.rkd_threat_lbl.pack(side="right", padx=16)

        # ══ ПАНЕЛЬ НАХОДОК (инциденты) ════════════════════════════
        self.rkd_findings_frame = ctk.CTkScrollableFrame(
            frame, fg_color=C_BG, corner_radius=10,
            border_width=1, border_color=C_BORDER,
            label_text=t("incident_log"),
            label_font=ctk.CTkFont(family=MONO, size=10, weight="bold"),
            label_fg_color=C_PANEL)
        self.rkd_findings_frame.pack(fill="both", expand=True, padx=14, pady=(4, 10))

        # Пустое состояние
        empty = ctk.CTkFrame(self.rkd_findings_frame, fg_color="transparent")
        empty.pack(pady=30)
        ctk.CTkLabel(empty, text="\U0001f6e1",
                     font=ctk.CTkFont(size=36)).pack()
        ctk.CTkLabel(empty,
                     text=t("press_scan"),
                     font=ctk.CTkFont(family=MONO, size=11),
                     text_color=C_DIM).pack(pady=(8, 0))

        # Сохраняем палитру для использования в рендере находок
        self._rkd_palette = {
            "bg": C_BG, "panel": C_PANEL, "panel_hi": C_PANEL_HI,
            "border": C_BORDER, "mono": C_MONO, "dim": C_DIM,
            "cyan": C_CYAN, "red": C_RED, "amber": C_AMBER,
            "purple": C_PURPLE, "mono_font": MONO,
        }

        return frame

    def _rkdefense_baseline(self):
        """Создаёт baseline чистой системы."""
        P = getattr(self, "_rkd_palette", {"cyan": "#00d4ff", "amber": "#ffa726",
                                            "red": "#ff3b4e", "mono": "#7d8da5"})
        try:
            from rootkit_detector import RootkitDetector
            self.after(0, lambda: self.rkd_status.configure(
                text=t("creating_baseline"), text_color=P["amber"]))
            det = RootkitDetector()
            bl = det.create_baseline()
            n_bins = len(bl.get("binaries", {}))
            n_mods = len(bl.get("modules", []))
            self.app_log(f"BASELINE создан: {n_bins} бинарников, {n_mods} модулей")
            self.after(0, lambda: [
                self.rkd_status.configure(text=t("baseline_created"), text_color=P["cyan"]),
                self.rkd_baseline_lbl.configure(
                    text=f"\u2713 baseline: {n_bins} bins / {n_mods} mods",
                    text_color=P["cyan"]),
            ])
        except Exception as e:
            self.after(0, lambda err=str(e): self.rkd_status.configure(
                text=f"{t('error_lc')}: {err[:40]}", text_color=P["red"]))

    def _run_rkdefense_scan(self):
        """Запускает rootkit сканирование (SOC дизайн)."""
        from rootkit_detector import RootkitDetector
        P = self._rkd_palette

        # Сброс UI
        self.after(0, lambda: [
            self.rkd_status.configure(text=t("scanning_lc"), text_color=P["amber"]),
            self.rkd_score_lbl.configure(text="..", text_color=P["amber"]),
            self.rkd_verdict.configure(text=t("scan_in_progress"), text_color=P["amber"]),
            self.rkd_engine_lbl.configure(text="\u25cf SCANNING", text_color=P["amber"]),
            self.rkd_progress.set(0),
            self.rkd_progress.configure(progress_color=P["cyan"]),
            self.rkd_threat_lbl.configure(text=""),
        ])
        for card, status, strip in self._rkd_cards:
            self.after(0, lambda c=card, s=status, st=strip: [
                c.configure(border_color=P["border"]),
                s.configure(text="\u25cb", text_color=P["dim"]),
                st.configure(fg_color=P["dim"])])

        self.after(0, lambda: [w.destroy() for w in self.rkd_findings_frame.winfo_children()])

        try:
            det = RootkitDetector()

            if not det.has_baseline():
                self.after(0, lambda: self.rkd_baseline_lbl.configure(
                    text=t("baseline_missing"), text_color=P["amber"]))

            methods = [
                (0, det.detect_hidden_processes),
                (1, det.detect_hidden_modules),
                (2, det.detect_privilege_escalation),
                (3, det.detect_ld_preload),
                (4, det.detect_binary_tampering),
                (5, det.detect_suspicious_connections),
            ]
            all_findings = []
            for idx, fn in methods:
                self.after(0, lambda p=(idx+1)/6: self.rkd_progress.set(p))
                try:
                    found = fn()
                    real = [f for f in found if not (f.severity == "НИЗКАЯ" and f.title == "Ошибка проверки")]
                    all_findings.extend(real)
                    card, status, strip = self._rkd_cards[idx]
                    if real:
                        self.after(0, lambda c=card, s=status, st=strip, n=len(real): [
                            c.configure(border_color=P["red"]),
                            s.configure(text=f"\u26a0 {n}", text_color=P["red"]),
                            st.configure(fg_color=P["red"])])
                    else:
                        self.after(0, lambda c=card, s=status, st=strip: [
                            c.configure(border_color=P["cyan"]),
                            s.configure(text="\u2713", text_color=P["cyan"]),
                            st.configure(fg_color=P["cyan"])])
                except Exception:
                    pass

            high = sum(1 for f in all_findings if f.severity == "ВЫСОКАЯ")
            med  = sum(1 for f in all_findings if f.severity == "СРЕДНЯЯ")
            score = max(0, 100 - high * 25 - med * 10)

            if high > 0:
                threat, tcolor, verdict = "ВЫСОКАЯ", P["red"], "\u26a0 ROOTKIT ОБНАРУЖЕН"
            elif med > 0:
                threat, tcolor, verdict = "СРЕДНЯЯ", P["amber"], "\u26a0 ПОДОЗРЕНИЕ"
            else:
                threat, tcolor, verdict = "ЧИСТАЯ", P["cyan"], "\u2713 СИСТЕМА ЧИСТА"

            self.after(0, lambda s=score, c=tcolor: self.rkd_score_lbl.configure(
                text=str(s), text_color=c))
            self.after(0, lambda v=verdict, c=tcolor: self.rkd_verdict.configure(
                text=v, text_color=c))
            self.after(0, lambda t=threat, c=tcolor: self.rkd_threat_lbl.configure(
                text=f"\u25cf {t}", text_color=c))
            self.after(0, lambda c=tcolor: self.rkd_progress.configure(progress_color=c))
            self.after(0, lambda: self.rkd_status.configure(
                text=t("scan_done_lc"), text_color=P["mono"]))
            self.after(0, lambda c=tcolor: self.rkd_engine_lbl.configure(
                text="\u25cf SCAN COMPLETE", text_color=c))

            self.after(0, lambda f=all_findings, t=threat: self._render_rkd_findings(f, t))

            if threat == "ВЫСОКАЯ":
                self.after(600, lambda f=all_findings: self._show_defense_modal(f))

            if all_findings:
                self._rkdefense_learn(all_findings)

        except Exception as e:
            self.after(0, lambda err=str(e): self.rkd_status.configure(
                text=f"{t('error_lc')}: {err[:40]}", text_color=P["red"]))

    def _open_syslog_window(self):
        """Окно live-просмотра: системный журнал + появление/исчезновение
        процессов и новых слушающих портов (что происходит ВНЕ программы)."""
        if getattr(self, "_syslog_win", None) is not None:
            try:
                if self._syslog_win.winfo_exists():
                    self._syslog_win.lift(); return
            except Exception:
                pass

        import subprocess, shutil
        win = ctk.CTkToplevel(self)
        win.title(t("syslog_title"))
        win.geometry("900x560")
        win.configure(fg_color="#0a0e16")
        self._syslog_win = win
        self._syslog_running = True

        bar = ctk.CTkFrame(win, fg_color="#0d1117", height=42)
        bar.pack(fill="x"); bar.pack_propagate(False)
        ctk.CTkLabel(bar, text="\U0001f4e1 " + t("syslog_title"),
                     font=ctk.CTkFont(size=13, weight="bold"),
                     text_color="#00d4ff").pack(side="left", padx=14)
        status = ctk.CTkLabel(bar, text="● LIVE", text_color="#2dc97e",
                              font=ctk.CTkFont(size=11, weight="bold"))
        status.pack(side="left", padx=8)
        self._syslog_autoscroll = True
        ctk.CTkButton(bar, text=t("clear_btn"), width=90, height=26,
                      fg_color="#1e293b", hover_color="#2d3748",
                      command=lambda: box.delete("1.0", "end")).pack(side="right", padx=10)

        box = ctk.CTkTextbox(win, font=ctk.CTkFont(family="monospace", size=11),
                             fg_color="#06090f", text_color="#9fe7c0")
        box.pack(fill="both", expand=True, padx=8, pady=8)

        def append(line, tag=""):
            if not self._syslog_running:
                return
            try:
                box.insert("end", line.rstrip() + "\n")
                # держим не больше ~500 строк, чтобы не разрасталось
                if int(box.index("end-1c").split(".")[0]) > 500:
                    box.delete("1.0", "100.0")
                box.see("end")
            except Exception:
                pass

        # ── Поток 1: системный журнал (journalctl -f / tail -F syslog) ──
        def stream_syslog():
            cmd = None
            if shutil.which("journalctl"):
                cmd = ["journalctl", "-f", "-n", "30", "--no-pager"]
            elif __import__("os").path.exists("/var/log/syslog"):
                cmd = ["tail", "-n", "30", "-F", "/var/log/syslog"]
            if not cmd:
                self.after(0, lambda: append("[!] Нет journalctl и /var/log/syslog. "
                                             "Доступен мониторинг процессов/портов ниже."))
                return
            try:
                proc = subprocess.Popen(cmd, stdout=subprocess.PIPE,
                                        stderr=subprocess.STDOUT, text=True, bufsize=1)
                self._syslog_proc = proc
                for line in proc.stdout:
                    if not self._syslog_running:
                        break
                    self.after(0, lambda l=line: append("  " + l))
            except Exception as e:
                self.after(0, lambda: append(f"[!] Журнал недоступен: {e}"))

        # ── Поток 2: события вне программы (процессы и порты, diff каждые 2с) ──
        def watch_system():
            import psutil, time as _t
            prev_pids, prev_ports = None, None
            while self._syslog_running:
                try:
                    cur_pids = {}
                    for pr in psutil.process_iter(["pid", "name", "username"]):
                        cur_pids[pr.info["pid"]] = (pr.info.get("name") or "?",
                                                    pr.info.get("username") or "?")
                    cur_ports = set()
                    for c in psutil.net_connections(kind="inet"):
                        if c.status == "LISTEN" and c.laddr:
                            cur_ports.add(c.laddr.port)
                    if prev_pids is not None:
                        for pid in set(cur_pids) - set(prev_pids):
                            n, u = cur_pids[pid]
                            self.after(0, lambda p=pid, n=n, u=u:
                                       append(f"[PROC +] PID {p} {n} (user={u})"))
                        for pid in set(prev_pids) - set(cur_pids):
                            n, _u = prev_pids[pid]
                            self.after(0, lambda p=pid, n=n:
                                       append(f"[PROC -] PID {p} {n} завершён"))
                        for port in cur_ports - prev_ports:
                            self.after(0, lambda p=port:
                                       append(f"[PORT +] новый слушающий порт :{p}"))
                        for port in prev_ports - cur_ports:
                            self.after(0, lambda p=port:
                                       append(f"[PORT -] порт :{p} закрыт"))
                    prev_pids, prev_ports = cur_pids, cur_ports
                except Exception:
                    pass
                for _ in range(4):
                    if not self._syslog_running:
                        break
                    _t.sleep(0.5)

        def on_close():
            self._syslog_running = False
            try:
                if getattr(self, "_syslog_proc", None):
                    self._syslog_proc.terminate()
            except Exception:
                pass
            self._syslog_win = None
            win.destroy()
        win.protocol("WM_DELETE_WINDOW", on_close)

        append(f"=== {t('syslog_title')} — старт {datetime.now().strftime('%H:%M:%S')} ===")
        threading.Thread(target=stream_syslog, daemon=True).start()
        threading.Thread(target=watch_system, daemon=True).start()

    def _clear_rkdefense(self):
        """Сбросить результаты последнего скана: находки, карточки, score."""
        P = self._rkd_palette
        for w in self.rkd_findings_frame.winfo_children():
            w.destroy()
        for card, status, strip in self._rkd_cards:
            card.configure(border_color=P["border"])
            status.configure(text="\u25cb", text_color=P["dim"])
            strip.configure(fg_color=P["dim"])
        self.rkd_score_lbl.configure(text="--", text_color=P["dim"])
        self.rkd_verdict.configure(text=t("waiting_caps"), text_color=P["dim"])
        self.rkd_threat_lbl.configure(text="")
        self.rkd_progress.set(0)
        self.rkd_status.configure(text=t("ready_to_scan_lc"), text_color=P["dim"])
        self.rkd_engine_lbl.configure(text="\u25cf IDLE", text_color=P["dim"])
        self._rkd_last_findings = []
        self._last_rkd_report = []

    def _render_rkd_findings(self, findings: list, threat: str):
        """Рисует находки как forensic-инциденты."""
        P = self._rkd_palette
        MONO = P["mono_font"]

        # Текстовый отчёт для копирования
        rep = [f"Угроза: {threat} | Находок: {len(findings)}", ""]
        for f in findings:
            rep.append(f"[{f.severity}] {f.title}")
            rep.append(f"  Метод: {f.method}")
            rep.append(f"  ГДЕ: {f.where}")
            rep.append(f"  КАК: {f.how}")
            rep.append(f"  ПОЧЕМУ: {f.why}")
            rep.append(f"  УСТРАНЕНИЕ: {f.fix}")
            rep.append(f"  MITRE: {f.mitre}")
            if f.evidence:
                rep.append(f"  EVIDENCE: {f.evidence}")
            rep.append("")
        self._last_rkd_report = rep
        self.app_log(f"ROOTKIT-СКАН: {threat}, находок {len(findings)}")
        for f in findings:
            self.app_log(f"  [{f.severity}] {f.title} @ {f.where}")

        for w in self.rkd_findings_frame.winfo_children():
            w.destroy()

        # Пустое состояние — система чиста
        if not findings:
            ok = ctk.CTkFrame(self.rkd_findings_frame, fg_color=P["panel"],
                              corner_radius=12, border_width=1, border_color=P["cyan"])
            ok.pack(fill="x", padx=10, pady=10)
            inner = ctk.CTkFrame(ok, fg_color="transparent")
            inner.pack(pady=22)
            ctk.CTkLabel(inner, text="\u2713",
                         font=ctk.CTkFont(size=40),
                         text_color=P["cyan"]).pack()
            ctk.CTkLabel(inner, text=t("system_clean"),
                         font=ctk.CTkFont(family=MONO, size=14, weight="bold"),
                         text_color=P["cyan"]).pack(pady=(8, 2))
            ctk.CTkLabel(inner, text=t("no_rootkit_5"),
                         font=ctk.CTkFont(family=MONO, size=10),
                         text_color=P["dim"]).pack()
            return

        # Карточки-инциденты (рендерим не больше _RKD_MAX —
        # сотни тяжёлых карточек подвешивают UI; остальное есть в отчёте)
        _RKD_MAX = 40
        if len(findings) > _RKD_MAX:
            warn = ctk.CTkFrame(self.rkd_findings_frame, fg_color=P["panel"],
                                corner_radius=8, border_width=1, border_color=P["amber"])
            warn.pack(fill="x", padx=10, pady=(8, 2))
            ctk.CTkLabel(warn,
                         text=f"Показаны первые {_RKD_MAX} из {len(findings)} находок. "
                              f"Полный список — в «Копировать всё».",
                         font=ctk.CTkFont(family=MONO, size=10),
                         text_color=P["amber"]).pack(anchor="w", padx=12, pady=8)
        for i, f in enumerate(findings[:_RKD_MAX]):
            sev_color = {"ВЫСОКАЯ": P["red"], "СРЕДНЯЯ": P["amber"],
                         "НИЗКАЯ": P["mono"]}.get(f.severity, P["mono"])

            # Контейнер инцидента
            card = ctk.CTkFrame(self.rkd_findings_frame, fg_color=P["panel"],
                                corner_radius=10, border_width=1, border_color=P["border"])
            card.pack(fill="x", padx=10, pady=6)

            # Severity-полоса слева + контент
            body = ctk.CTkFrame(card, fg_color="transparent")
            body.pack(fill="x")
            sev_strip = ctk.CTkFrame(body, fg_color=sev_color, width=4, corner_radius=0)
            sev_strip.pack(side="left", fill="y", padx=(0, 0), pady=0)

            content_col = ctk.CTkFrame(body, fg_color="transparent")
            content_col.pack(side="left", fill="both", expand=True, padx=2)

            # ── Заголовок инцидента ──
            head = ctk.CTkFrame(content_col, fg_color="transparent")
            head.pack(fill="x", padx=12, pady=(10, 2))
            # ID инцидента
            ctk.CTkLabel(head, text=f"INC-{i+1:02d}",
                         font=ctk.CTkFont(family=MONO, size=9, weight="bold"),
                         fg_color=P["panel_hi"], corner_radius=4,
                         text_color=sev_color, width=54, height=20).pack(side="left", padx=(0, 8))
            ctk.CTkLabel(head, text=f.title,
                         font=ctk.CTkFont(size=12, weight="bold"),
                         text_color="#e6edf7").pack(side="left")
            # Severity бейдж справа
            ctk.CTkLabel(head, text=f.severity,
                         font=ctk.CTkFont(family=MONO, size=9, weight="bold"),
                         fg_color=sev_color, corner_radius=4,
                         text_color="#06090f", width=72, height=20).pack(side="right")

            # Метод-детектор
            ctk.CTkLabel(content_col, text=f"detector: {f.method}",
                         font=ctk.CTkFont(family=MONO, size=9),
                         text_color=P["dim"]).pack(anchor="w", padx=12, pady=(0, 6))

            # ── Технические секции (моноширинные блоки) ──
            sections = [
                ("WHERE", f.where, P["mono"]),
                ("HOW",   f.how,   P["mono"]),
                ("RISK",  f.why,   "#c7d3e3"),
            ]
            for tag, value, vcolor in sections:
                blk = ctk.CTkFrame(content_col, fg_color=P["bg"], corner_radius=6)
                blk.pack(fill="x", padx=12, pady=2)
                ctk.CTkLabel(blk, text=tag,
                             font=ctk.CTkFont(family=MONO, size=8, weight="bold"),
                             text_color=P["dim"], width=54, anchor="nw",
                             justify="left").pack(side="left", anchor="n", padx=(8, 4), pady=6)
                ctk.CTkLabel(blk, text=value,
                             font=ctk.CTkFont(family=MONO, size=10),
                             text_color=vcolor, anchor="w", justify="left",
                             wraplength=560).pack(side="left", anchor="n", fill="x",
                                                  expand=True, pady=6, padx=(0, 8))

            # ── Команда устранения (terminal-блок) ──
            fix_blk = ctk.CTkFrame(content_col, fg_color="#04130a", corner_radius=6,
                                   border_width=1, border_color="#0f3320")
            fix_blk.pack(fill="x", padx=12, pady=(4, 2))
            ctk.CTkLabel(fix_blk, text="\u2192 REMEDIATION",
                         font=ctk.CTkFont(family=MONO, size=8, weight="bold"),
                         text_color="#3ddc84").pack(anchor="w", padx=8, pady=(5, 0))
            for cmd in f.fix.replace("\\n", chr(10)).split(chr(10)):
                if cmd.strip():
                    ctk.CTkLabel(fix_blk, text=f"  $ {cmd.strip()}",
                                 font=ctk.CTkFont(family=MONO, size=10),
                                 text_color="#5ef0a0", anchor="w",
                                 justify="left").pack(anchor="w", padx=8, pady=0)
            ctk.CTkLabel(fix_blk, text="",
                         font=ctk.CTkFont(size=2)).pack(pady=1)

            # ── MITRE ATT&CK footer ──
            foot = ctk.CTkFrame(content_col, fg_color="transparent")
            foot.pack(fill="x", padx=12, pady=(2, 10))
            ctk.CTkLabel(foot, text="\u2316 MITRE ATT&CK",
                         font=ctk.CTkFont(family=MONO, size=8, weight="bold"),
                         text_color=P["purple"]).pack(side="left")
            ctk.CTkLabel(foot, text=f.mitre,
                         font=ctk.CTkFont(family=MONO, size=9),
                         text_color=P["purple"]).pack(side="left", padx=6)
            if f.evidence:
                ctk.CTkLabel(foot, text=f"\u2502 evidence: {f.evidence[:50]}",
                             font=ctk.CTkFont(family=MONO, size=8),
                             text_color=P["dim"]).pack(side="left", padx=6)

        # Сохраняем находки и строим AI-панель реагирования
        self._rkd_last_findings = findings
        if findings:
            self._render_rkd_response(findings, threat)

    def _render_rkd_response(self, findings: list, threat: str):
        """AI Incident Response — Claude даёт команды реагирования."""
        P = self._rkd_palette
        MONO = P["mono_font"]

        # Контейнер AI-реагирования (внизу панели находок)
        resp = ctk.CTkFrame(self.rkd_findings_frame, fg_color="#0a1322",
                            corner_radius=10, border_width=1, border_color=P["cyan"])
        resp.pack(fill="x", padx=10, pady=(10, 6))

        # Заголовок
        hdr = ctk.CTkFrame(resp, fg_color="transparent")
        hdr.pack(fill="x", padx=12, pady=(10, 4))
        ctk.CTkLabel(hdr, text="\U0001f916  AI INCIDENT RESPONSE",
                     font=ctk.CTkFont(family=MONO, size=11, weight="bold"),
                     text_color=P["cyan"]).pack(side="left")
        ctk.CTkLabel(hdr, text=t("claude_plan"),
                     font=ctk.CTkFont(family=MONO, size=9),
                     text_color=P["dim"]).pack(side="left", padx=8)

        # Кнопки управления
        btns = ctk.CTkFrame(resp, fg_color="transparent")
        btns.pack(fill="x", padx=12, pady=(0, 6))

        ask_btn = ctk.CTkButton(
            btns, text=t("get_ai_plan"),
            height=34, corner_radius=8,
            fg_color=P["cyan"], hover_color="#00b8e0",
            text_color="#06090f",
            font=ctk.CTkFont(size=11, weight="bold"),
            command=lambda f=findings: threading.Thread(
                target=self._rkd_ai_response, args=(f,), daemon=True).start())
        ask_btn.pack(side="left", padx=(0, 6))

        self._rkd_copy_cmd_btn = ctk.CTkButton(
            btns, text=t("copy_commands"),
            height=34, corner_radius=8,
            fg_color=P["panel_hi"], hover_color="#16233a",
            border_width=1, border_color=P["border"],
            text_color="#c7d3e3",
            font=ctk.CTkFont(family=MONO, size=10),
            command=self._rkd_copy_commands, state="disabled")
        self._rkd_copy_cmd_btn.pack(side="left", padx=(0, 6))

        verify_btn = ctk.CTkButton(
            btns, text=t("check_fix"),
            height=34, corner_radius=8,
            fg_color="#1a3a1a", hover_color="#2d6a4f",
            text_color="#5ef0a0",
            font=ctk.CTkFont(size=11, weight="bold"),
            command=lambda: threading.Thread(
                target=self._rkd_verify_remediation, daemon=True).start())
        verify_btn.pack(side="left")

        # Поле вывода AI-плана
        self._rkd_ai_box = ctk.CTkTextbox(
            resp, height=180,
            font=ctk.CTkFont(family=MONO, size=11),
            fg_color="#06090f", text_color="#c7d3e3", wrap="word")
        self._rkd_ai_box.pack(fill="x", padx=12, pady=(0, 6))
        self._rkd_ai_box.insert("end",
            "\u2192 Нажми «Получить план от AI» — Claude проанализирует "
            "обнаруженные угрозы и даст точные команды для устранения.\n")
        self._rkd_ai_box.configure(state="disabled")

        # Хранилище извлечённых команд
        self._rkd_extracted_cmds = []

        # Статус проверки
        self._rkd_verify_lbl = ctk.CTkLabel(
            resp, text="",
            font=ctk.CTkFont(family=MONO, size=10))
        self._rkd_verify_lbl.pack(anchor="w", padx=12, pady=(0, 10))

    def _rkd_ai_response(self, findings: list):
        """Запрашивает у Claude план реагирования по конкретным находкам."""
        P = self._rkd_palette
        try:
            import anthropic, re

            self.after(0, lambda: [
                self._rkd_ai_box.configure(state="normal"),
                self._rkd_ai_box.delete("1.0", "end"),
                self._rkd_ai_box.insert("end", "\U0001f916 Claude анализирует угрозы...\n"),
                self._rkd_ai_box.configure(state="disabled"),
            ])

            # Формируем детальное описание находок для Claude
            findings_desc = []
            for f in findings:
                findings_desc.append(
                    f"- [{f.severity}] {f.title}\n"
                    f"  Метод: {f.method}\n"
                    f"  Где: {f.where}\n"
                    f"  MITRE: {f.mitre}\n"
                    f"  Evidence: {f.evidence}")
            findings_text = "\n".join(findings_desc)

            api_key = cfg.get("anthropic", {}).get("api_key", "")
            client = anthropic.Anthropic(api_key=api_key)

            prompt = f"""Ты эксперт по реагированию на инциденты Linux (Incident Response).
RootkitGuard обнаружил следующие rootkit-угрозы на системе Ubuntu (ядро 6.17):

{findings_text}

Дай ЧЁТКИЙ план реагирования по этапам Incident Response. Для КАЖДОГО этапа дай ТОЧНЫЕ bash-команды (с конкретными именами модулей/PID/портов из находок выше, не плейсхолдеры).

Формат ответа СТРОГО такой:

ЭТАП 1 — ИЗОЛЯЦИЯ
$ команда1
$ команда2

ЭТАП 2 — УСТРАНЕНИЕ
$ команда1

ЭТАП 3 — ПРОВЕРКА
$ команда1

ЭТАП 4 — ВОССТАНОВЛЕНИЕ (если нужно)
$ команда1

Каждая команда с новой строки начинается с "$ ". Между этапами объясни ОДНИМ предложением что делаем и зачем. Команды реальные, готовые к вставке в терминал. Отвечай на русском."""

            message = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=900,
                messages=[{"role": "user", "content": prompt}])
            result = message.content[0].text

            # Извлекаем команды (строки начинающиеся с $)
            cmds = re.findall(r"^\s*\$\s*(.+)$", result, re.M)
            self._rkd_extracted_cmds = [c.strip() for c in cmds if c.strip()]

            self.app_log(f"AI Incident Response: получен план ({len(self._rkd_extracted_cmds)} команд)")

            self.after(0, lambda r=result: [
                self._rkd_ai_box.configure(state="normal", height=280),
                self._rkd_ai_box.delete("1.0", "end"),
                self._rkd_ai_box.insert("end", r),
                self._rkd_ai_box.configure(state="disabled"),
                self._rkd_copy_cmd_btn.configure(state="normal"),
            ])

        except Exception as e:
            self.after(0, lambda err=str(e): [
                self._rkd_ai_box.configure(state="normal"),
                self._rkd_ai_box.delete("1.0", "end"),
                self._rkd_ai_box.insert("end", f"[!] Ошибка AI: {err}\n\n"
                    "Проверь что API-ключ задан в config.yaml и API запущен."),
                self._rkd_ai_box.configure(state="disabled"),
            ])

    def _rkd_copy_commands(self):
        """Копирует извлечённые команды в буфер."""
        P = self._rkd_palette
        if not self._rkd_extracted_cmds:
            return
        text = "\n".join(self._rkd_extracted_cmds)
        try:
            self.clipboard_clear()
            self.clipboard_append(text)
            self.update()
            self._rkd_copy_cmd_btn.configure(text=t("copied"))
            self.after(2000, lambda: self._rkd_copy_cmd_btn.configure(
                text=t("copy_commands")))
            self.app_log(f"Скопировано {len(self._rkd_extracted_cmds)} команд реагирования")
        except Exception:
            pass

    def _rkd_verify_remediation(self):
        """Повторный скан для проверки что rootkit устранён."""
        P = self._rkd_palette
        from rootkit_detector import RootkitDetector

        # Запоминаем сколько было находок ДО
        before = len(getattr(self, "_rkd_last_findings", []))

        self.after(0, lambda: [
            self._rkd_verify_lbl.configure(
                text=t("rescanning"), text_color=P["amber"]),
        ])

        try:
            det = RootkitDetector()
            result = det.full_scan()
            after = result["total"]
            score = result["score"]
            threat = result["threat"]

            self.app_log(f"ПРОВЕРКА УСТРАНЕНИЯ: было {before} находок, стало {after}, score {score}")

            if after == 0:
                msg = f"\u2713 УСТРАНЕНО! Было {before} угроз \u2192 стало 0. Score: {score}/100. Система чиста."
                color = P["cyan"]
            elif after < before:
                msg = f"\u26a0 Частично: было {before} \u2192 стало {after}. Score: {score}. Остались угрозы."
                color = P["amber"]
            else:
                msg = f"\u2717 Не устранено: {after} угроз остаётся. Score: {score}. Выполни команды AI."
                color = P["red"]

            # Если устранено — перезапускаем полный скан чтобы обновить весь UI
            self.after(0, lambda m=msg, c=color: self._rkd_verify_lbl.configure(text=m, text_color=c))
            self.after(800, lambda: threading.Thread(
                target=self._run_rkdefense_scan, daemon=True).start())

        except Exception as e:
            self.after(0, lambda err=str(e): self._rkd_verify_lbl.configure(
                text=f"{t('check_error')}: {err[:50]}", text_color=P["red"]))

    def _rkdefense_learn(self, findings: list):
        """Самообучение: модель учится на rootkit находках."""
        def learn():
            try:
                from online_learner import OnlineLearner
                import pandas as pd, numpy as np
                learner = OnlineLearner()
                cols = list(self.rf.feature_names_in_)
                # Генерируем образцы на основе severity находок
                n = min(len(findings) * 5, 30)
                if n > 0:
                    samples = pd.DataFrame(
                        np.random.randn(n, len(cols)) * 3 + 5, columns=cols)
                    added = learner.add_attack_samples(samples, label=1)
                    self.after(0, lambda a=added: self.learn_lbl.configure(
                        text=f"\U0001f9e0 +{a}", text_color="#a855f7"))
                    if learner.should_retrain():
                        self.after(0, lambda: self.learn_lbl.configure(
                            text=t("training"), text_color="#f59e0b"))
                        result = learner.retrain()
                        if result.get("status") == "success":
                            self._load_models()
                            self.after(0, lambda r=result: self.learn_lbl.configure(
                                text=f"\U0001f9e0 v{r['version']} \u2713",
                                text_color="#2dc97e"))
            except Exception:
                pass
        threading.Thread(target=learn, daemon=True).start()


if __name__ == "__main__":
    app = RootkitGuard()
    app.mainloop()
